import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_font_for_all(doc, font_name='Times New Roman', font_size=12):
    """设置所有段落的字体"""
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)

def add_line_numbers(doc):
    """添加行号"""
    for section in doc.sections:
        sectPr = section._sectPr
        lnNumType = sectPr.makeelement(qn('w:lnNumType'), {
            qn('w:countBy'): '1',
            qn('w:restart'): 'newPage'
        })
        sectPr.append(lnNumType)

def add_page_numbers(doc):
    """添加页码"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        
        # 添加页码
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加页码字段
        run = paragraph.add_run()
        fldChar1 = docx.oxml.OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._element.append(fldChar1)
        
        run2 = paragraph.add_run()
        instrText = docx.oxml.OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        run2._element.append(instrText)
        
        run3 = paragraph.add_run()
        fldChar2 = docx.oxml.OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run3._element.append(fldChar2)

# 加载文档
doc = docx.Document("/Volumes/agents/codex/paper2-LSTM量化金融/archive_20260710_1046_chatgpt_review/JFDS_Submission_Package_20260712/paper/论文_公式恢复与图表美化.docx")

print("=== 修复排版 ===\n")

# 1. 修改字体
print("1. 修改字体为 Times New Roman 12pt...")
set_font_for_all(doc, 'Times New Roman', 12)
print("   ✅ 完成")

# 2. 添加行号
print("\n2. 添加行号...")
add_line_numbers(doc)
print("   ✅ 完成")

# 3. 添加页码
print("\n3. 添加页码...")
add_page_numbers(doc)
print("   ✅ 完成")

# 4. 保存文档
output_file = "/Volumes/agents/codex/paper2-LSTM量化金融/archive_20260710_1046_chatgpt_review/JFDS_Submission_Package_20260712/paper/论文_公式恢复与图表美化_FINAL.docx"
doc.save(output_file)

print(f"\n✅ 排版修复完成！")
print(f"保存为: {output_file}")

# 5. 验证修改
print("\n=== 验证修改 ===")
doc2 = docx.Document(output_file)

fonts = set()
for para in doc2.paragraphs:
    for run in para.runs:
        if run.font.name:
            fonts.add(run.font.name)

print(f"字体: {fonts}")

# 检查行号
has_line_numbers = False
for section in doc2.sections:
    sectPr = section._sectPr
    for child in sectPr:
        if 'lnNumType' in child.tag:
            has_line_numbers = True
            break

print(f"行号: {'有' if has_line_numbers else '无'}")

# 检查页码
has_page_numbers = False
for section in doc2.sections:
    footer = section.footer
    if footer.paragraphs:
        for para in footer.paragraphs:
            if para.text.strip():
                has_page_numbers = True
                break

print(f"页码: {'有' if has_page_numbers else '无'}")

