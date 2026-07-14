import docx
from docx.shared import Pt
import re

# Load the document
doc = docx.Document("main_paper_jfds_READY_v15_logic_fixed.docx")

print("=== 清理稿件 ===\n")

# 1. 删除YAML元数据（段落4-7）
print("1. 删除YAML元数据...")
paragraphs_to_delete = []
for i in range(4, 8):
    if i < len(doc.paragraphs):
        text = doc.paragraphs[i].text.strip()
        if any(text.startswith(p) for p in ['title:', 'author:', 'date:', 'journal:']):
            paragraphs_to_delete.append(i)
            print(f"   删除段落 {i}: {text}")

# 删除段落（从后往前删除以避免索引问题）
for i in sorted(paragraphs_to_delete, reverse=True):
    p = doc.paragraphs[i]
    p._element.getparent().remove(p._element)

# 2. 修复Acknowledgments中的中文
print("\n2. 修复Acknowledgments中的中文...")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if 'Acknowledgments' in text:
        # 找到下一段落并替换中文
        if i+1 < len(doc.paragraphs):
            next_para = doc.paragraphs[i+1]
            next_text = next_para.text.strip()
            if any('\u4e00' <= char <= '\u9fff' for char in next_text):
                # 替换为英文
                new_text = "The author thanks the editor and anonymous reviewers for their valuable suggestions."
                next_para.clear()
                next_para.add_run(new_text)
                print(f"   修复段落 {i+1}: {next_text[:40]}... -> {new_text}")

# 3. 删除JSON代码块
print("\n3. 删除JSON代码块...")
json_start = None
json_end = None
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '"experiment": {' in text:
        json_start = i
    if json_start and i > json_start and text == '}':
        json_end = i
        break

if json_start and json_end:
    for i in range(json_start, json_end + 1):
        if i < len(doc.paragraphs):
            p = doc.paragraphs[i]
            p._element.getparent().remove(p._element)
            print(f"   删除段落 {i}")

# 4. 删除营销式表达
print("\n4. 删除营销式表达...")
marketing_replacements = {
    'implemented by any quantitative team in under one day': 'implementable by quantitative teams',
    'saves hundreds of thousands of dollars in annual transaction costs': 'reduces transaction costs',
    'immediate practical impact': 'practical implications',
    'actionable value': 'practical value',
}

for i, para in enumerate(doc.paragraphs):
    original_text = para.text
    fixed_text = original_text
    for old, new in marketing_replacements.items():
        if old in fixed_text.lower():
            pattern = re.compile(re.escape(old), re.IGNORECASE)
            fixed_text = pattern.sub(new, fixed_text)
    
    if fixed_text != original_text:
        para.clear()
        para.add_run(fixed_text)
        print(f"   修复段落 {i}")

# 5. 保存文档
output_file = "main_paper_jfds_READY_v16_cleaned.docx"
doc.save(output_file)

print(f"\n✅ 稿件清理完成！")
print(f"保存为: {output_file}")

