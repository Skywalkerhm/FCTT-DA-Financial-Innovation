import docx
from docx.shared import Pt
from docx.oxml.ns import qn

def read_markdown_file(filepath):
    """Read a markdown file and return its content."""
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()

def add_appendix_to_doc(doc, markdown_content):
    """Add appendix content to the document."""
    lines = markdown_content.strip().split('\n')
    
    # Find the last paragraph in the document
    last_para = doc.paragraphs[-1]
    
    # Add content after the last paragraph
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Determine style based on markdown formatting
        if line.startswith('## '):
            # Main heading
            para = doc.add_paragraph(line[3:], style='Heading 1')
            last_para._element.addnext(para._element)
            last_para = para
        elif line.startswith('### '):
            # Sub heading
            para = doc.add_paragraph(line[4:], style='Heading 2')
            last_para._element.addnext(para._element)
            last_para = para
        elif line.startswith('**') and line.endswith('**'):
            # Bold text
            para = doc.add_paragraph(line[2:-2])
            if para.runs:
                para.runs[0].bold = True
            last_para._element.addnext(para._element)
            last_para = para
        elif line.startswith('**Table'):
            # Table caption
            para = doc.add_paragraph(line)
            if para.runs:
                para.runs[0].bold = True
            last_para._element.addnext(para._element)
            last_para = para
        elif line.startswith('- '):
            # Bullet point
            para = doc.add_paragraph(line[2:], style='List Bullet')
            last_para._element.addnext(para._element)
            last_para = para
        elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
            # Numbered list
            para = doc.add_paragraph(line, style='List Number')
            last_para._element.addnext(para._element)
            last_para = para
        else:
            # Regular paragraph
            para = doc.add_paragraph(line)
            last_para._element.addnext(para._element)
            last_para = para

# Load the document
doc = docx.Document("main_paper_jfds_READY_v6_calibration_fixed.docx")

print("=== 添加α参数敏感性分析附录 ===\n")

# Read the appendix content
appendix_content = read_markdown_file("appendix_h_alpha_sensitivity.md")

# Add to document
add_appendix_to_doc(doc, appendix_content)

# Save the document
output_file = "main_paper_jfds_READY_v7_alpha_sensitivity.docx"
doc.save(output_file)

print("✅ α参数敏感性分析附录添加完成！")
print(f"保存为: {output_file}")

# Count total words
import re
total_words = sum(len(re.findall(r'\w+', para.text)) for para in doc.paragraphs)
print(f"\n文档总字数: {total_words}")

