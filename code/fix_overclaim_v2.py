import docx
from docx.shared import Pt
import re

# 加载文档
doc = docx.Document("/Volumes/agents/codex/paper2-LSTM量化金融/archive_20260710_1046_chatgpt_review/JFDS_Submission_Package_20260712/paper/论文_公式恢复与图表美化_NOGATING_FIXED.docx")

print("=== 修正过强表述 ===\n")

# 定义修正规则
replacements = [
    # 1. 修正 "algorithmic-agnostic"
    ('algorithmic-agnostic', 'persists across the four policy classes examined'),
    ('algorithm-agnostic', 'persists across the four policy classes examined'),
    ('Algorithm-agnostic', 'Persists across the four policy classes examined'),
    ('algorithm agnostic', 'persists across the four policy classes examined'),
    
    # 2. 修正过强的结论
    ('The failure mode is algorithmic-agnostic', 
     'The failure mode persists across the four policy classes examined (LinUCB, KernelUCB, NeuralUCB, Random Forest UCB)'),
    
    # 3. 修正关于context的过强表述
    ('context has value', 'context has theoretical value based on Oracle analysis'),
    ('context has incremental value', 'context has theoretical incremental value'),
    
    # 4. 修正behavioral diversity的表述
    ('behavioral diversity', 'action diversity'),
    ('behavioral adaptation', 'action variation'),
    
    # 5. 修正关于collapse原因的表述
    ('collapse is driven by source-side degeneracy', 
     'collapse is associated with source-side degeneracy, though the causal mechanism requires further investigation'),
    
    # 6. 修正关于random context的解释
    ('This confirms that collapse is driven by source-side degeneracy, not the informativeness of context features',
     'This suggests that collapse may be driven by the frozen policy parameters rather than context informativeness, raising questions about whether the contextual bandit effectively utilizes context'),
]

# 修正文档
fixed_count = 0
for i, para in enumerate(doc.paragraphs):
    original_text = para.text
    fixed_text = original_text
    
    # 应用替换
    for old, new in replacements:
        if old in fixed_text:
            fixed_text = fixed_text.replace(old, new)
            fixed_count += 1
    
    # 更新段落
    if fixed_text != original_text:
        para.clear()
        para.add_run(fixed_text)
        print(f"  修正段落 {i}:")
        print(f"    原文: {original_text[:80]}...")
        print(f"    修正: {fixed_text[:80]}...")

# 保存文档
output_file = "/Volumes/agents/codex/paper2-LSTM量化金融/archive_20260710_1046_chatgpt_review/JFDS_Submission_Package_20260712/paper/论文_公式恢复与图表美化_OVERCLAIM_FIXED.docx"
doc.save(output_file)

print(f"\n✅ 过强表述修正完成！")
print(f"修正数量: {fixed_count}")
print(f"保存为: {output_file}")

