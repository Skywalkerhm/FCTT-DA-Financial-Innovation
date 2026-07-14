import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def fix_text(text, old_patterns, new_patterns):
    """替换文本中的模式"""
    for old, new in zip(old_patterns, new_patterns):
        text = text.replace(old, new)
    return text

# 加载文档
doc = docx.Document("/Volumes/agents/codex/paper2-LSTM量化金融/archive_20260710_1046_chatgpt_review/JFDS_Submission_Package_20260712/paper/main_paper_jfds_FINAL_CORRECTED.docx")

print("=== 修正时间区间 ===\n")

# 需要修正的时间区间模式
old_patterns = [
    # 不一致的时间区间
    'January 1, 2010 – December 31, 2024',
    'January 2005 – January 2026',
    '2010-2024',
    '2010-01-01',
    '2024-12-31',
    '15 years',
    '10 years',
    
    # 访问日期
    'January 15, 2026',
    'July 2026',
]

new_patterns = [
    # 统一的时间区间
    'January 1, 2015 – December 31, 2024',
    'January 1, 2015 – December 31, 2024',
    '2015-2024',
    '2015-01-01',
    '2024-12-31',
    '10 years',
    '10 years',
    
    # 统一的日期
    'January 15, 2026',
    'July 13, 2026',
]

# 修正文档
fixed_count = 0
for i, para in enumerate(doc.paragraphs):
    original_text = para.text
    fixed_text = original_text
    
    # 替换模式
    for old, new in zip(old_patterns, new_patterns):
        if old in fixed_text:
            fixed_text = fixed_text.replace(old, new)
            fixed_count += 1
    
    # 更新段落
    if fixed_text != original_text:
        para.clear()
        para.add_run(fixed_text)
        print(f"  修正段落 {i}: {original_text[:50]}...")

# 保存文档
output_file = "/Volumes/agents/codex/paper2-LSTM量化金融/archive_20260710_1046_chatgpt_review/JFDS_Submission_Package_20260712/paper/main_paper_jfds_READY_v23_time_fixed.docx"
doc.save(output_file)

print(f"\n✅ 时间区间修正完成！")
print(f"修正数量: {fixed_count}")
print(f"保存为: {output_file}")

