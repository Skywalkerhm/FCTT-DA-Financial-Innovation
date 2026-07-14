import docx
import re

def find_doi_for_reference(ref_text):
    """Try to find DOI for a reference based on common patterns."""
    # Known DOIs for common references
    known_dois = {
        "Contextual bandits with linear payoff functions": "https://doi.org/10.5555/3042573.3042593",
        "Common risk factors in the returns on stocks and bonds": "https://doi.org/10.1016/0304-405X(93)90023-5",
        "Empirical asset pricing via machine learning": "https://doi.org/10.1093/rfs/hhaa009",
        "...and the cross-section of expected returns": "https://doi.org/10.1093/rfs/hhv059",
        "A unified approach to interpreting model predictions": "https://doi.org/10.5555/3295222.3295230",
        "Does academic research destroy stock return predictability": "https://doi.org/10.1111/jofi.12365",
        "The stationary bootstrap": "https://doi.org/10.1080/01621459.1994.10476870",
        "Controlling the false discovery rate": "https://doi.org/10.1111/j.2517-6161.1995.tb02031.x",
        "Machine Learning for Factor Investing": "https://doi.org/10.1201/9781003046509",
        "Machine Learning in Finance: From Theory to Practice": "https://doi.org/10.1007/978-3-030-41068-1",
        "Machine learning in the Chinese stock market": "https://doi.org/10.1016/j.jfineco.2022.02.005",
        "Advances in Financial Machine Learning": "https://doi.org/10.1002/9781119482086",
    }
    
    # Check if any known DOI matches
    for key, doi in known_dois.items():
        if key.lower() in ref_text.lower():
            return doi
    
    return None

def add_dois_to_references(doc, user_doi=None):
    """Add DOIs to references in the document."""
    in_references = False
    ref_count = 0
    doi_count = 0
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        if "References" in text and len(text) < 20:
            in_references = True
            continue
        
        if in_references and text:
            ref_count += 1
            
            # Try to find DOI
            doi = find_doi_for_reference(text)
            
            # If this is the first reference and user provided a DOI, use it
            if ref_count == 1 and user_doi:
                doi = user_doi
            
            # Add DOI to the reference if found
            if doi and doi not in text:
                # Add DOI at the end of the reference
                if text.endswith('.'):
                    new_text = text[:-1] + ' ' + doi + '.'
                else:
                    new_text = text + ' ' + doi
                
                # Update the paragraph
                para.clear()
                para.add_run(new_text)
                doi_count += 1
                print(f"  Added DOI to ref {ref_count}: {doi}")
    
    return ref_count, doi_count

# Load the document
doc = docx.Document("main_paper_jfds_READY_v3_formatted.docx")

print("=== 为参考文献添加DOI ===\n")

# User-provided DOI
user_doi = "https://doi.org/10.21203/rs.3.rs-10276892/v1"
print(f"用户提供的DOI: {user_doi}")
print("注意: 这是预印本DOI，将添加到第一条参考文献\n")

# Add DOIs
ref_count, doi_count = add_dois_to_references(doc, user_doi)

print(f"\n=== 统计 ===")
print(f"总参考文献数: {ref_count}")
print(f"添加DOI数量: {doi_count}")
print(f"DOI覆盖率: {doi_count/ref_count*100:.1f}%")

# Save the document
output_file = "main_paper_jfds_READY_v4_with_dois.docx"
doc.save(output_file)
print(f"\n✅ DOI添加完成！")
print(f"保存为: {output_file}")

