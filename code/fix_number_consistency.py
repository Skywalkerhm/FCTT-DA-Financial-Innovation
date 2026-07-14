import docx
from docx.shared import Pt
import re

# 加载文档
doc = docx.Document("/Volumes/agents/codex/paper2-LSTM量化金融/archive_20260710_1046_chatgpt_review/JFDS_Submission_Package_20260712/paper/论文_公式恢复与图表美化_RHETORIC.docx")

print("=== 修复数字一致性问题 ===\n")

# 核心问题：目标资产包含源资产，排除自身后是38个，总对数是228
replacements = [
    # 234 -> 228
    ('234 source-target pairs', '228 source-target pairs'),
    ('234 source–target pairs', '228 source–target pairs'),
    ('Across 234 source-target pairs', 'Across 228 source-target pairs'),
    ('Across 234 source–target pairs', 'Across 228 source–target pairs'),
    ('of 234 pairs', 'of 228 pairs'),
    ('of 234 source-target pairs', 'of 228 source-target pairs'),
    ('of 234 source–target pairs', 'of 228 source–target pairs'),
    ('100% of 234 pairs', '100% of 228 pairs'),
    ('234 total source-target pairs', '228 total source-target pairs'),
    ('total 234', 'total 228'),
    ('N_pairs = 234', 'N_pairs = 228'),
    ('N_{\\text{pairs}} = 234', 'N_{\\text{pairs}} = 228'),
    
    # 6 × 39 = 234 -> 6 × 38 = 228
    ('6×39=234', '6×38=228'),
    ('6 sources × 39 targets', '6 sources × 38 targets (excluding self)'),
    ('6 × 39 = 234', '6 × 38 = 228'),
    ('6 source × 39 target', '6 source × 38 target (excluding self)'),
    
    # BH-FDR for 234 pairs -> 228
    ('BH-FDR at 5%', 'BH-FDR at 5%'),
    ('234 pairs', '228 pairs'),
]

# 修正文档
fixed_count = 0
for i, para in enumerate(doc.paragraphs):
    original_text = para.text
    fixed_text = original_text
    
    # 替换模式
    for old, new in replacements:
        if old in fixed_text:
            fixed_text = fixed_text.replace(old, new)
            fixed_count += 1
    
    # 更新段落
    if fixed_text != original_text:
        para.clear()
        para.add_run(fixed_text)
        print(f"  修正段落 {i}: {original_text[:60]}...")

# 保存文档
output_file = "/Volumes/agents/codex/paper2-LSTM量化金融/archive_20260710_1046_chatgpt_review/JFDS_Submission_Package_20260712/paper/论文_公式恢复与图表美化_NUMBER_FIXED.docx"
doc.save(output_file)

print(f"\n✅ 数字一致性修复完成！")
print(f"修正数量: {fixed_count}")
print(f"保存为: {output_file}")

