import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color
    })
    shading_elm.append(shading)

def add_dataframe_to_doc(doc, df, caption, font_size=9):
    """将DataFrame添加到文档中"""
    # 添加表格标题
    p = doc.add_paragraph()
    run = p.add_run(caption)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 创建表格
    rows, cols = df.shape
    table = doc.add_table(rows=rows+1, cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 添加表头
    header = table.rows[0]
    for i, col_name in enumerate(df.columns):
        cell = header.cells[i]
        cell.text = str(col_name)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(font_size)
                run.font.name = 'Times New Roman'
        set_cell_shading(cell, 'D9E2F3')  # 浅蓝色背景
    
    # 添加数据行
    for i, (_, row) in enumerate(df.iterrows()):
        for j, value in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
                    run.font.name = 'Times New Roman'
    
    # 添加空行
    doc.add_paragraph()
    
    return table

def main():
    # 加载表格数据
    data_dir = "/Volumes/agents/codex/paper2-LSTM量化金融/archive_20260710_1046_chatgpt_review/JFDS_Submission_Package_20260712/paper/table_data"
    
    # 创建新文档
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # 添加标题
    title = doc.add_heading('Tables with Complete Data', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('This document contains all tables with complete numerical data for the paper:')
    doc.add_paragraph('"Parameters Transfer, Policies May Not: Diagnosing Static-Rule Collapse in Cross-Asset Decision Gating"')
    doc.add_paragraph()
    
    # 加载并添加所有表格
    table_configs = [
        ('Table1_Source_Diversity.csv', 'Table 1: Source Policy Diversity'),
        ('Table2_Collapse_by_Source.csv', 'Table 2: Collapse Rate by Source Asset'),
        ('Table3_TLT_Performance.csv', 'Table 3: TLT Source Policy Performance'),
        ('Table4_Collapse_vs_Threshold.csv', 'Table 4: Collapse Rate vs. Modal Share Threshold'),
        ('Table5_Performance_by_Class.csv', 'Table 5: Performance by Asset Class (TLT Source)'),
        ('Table6_Transaction_Costs.csv', 'Table 6: Transaction Cost and Turnover Summary'),
        ('Table7_Calibration_Ablation.csv', 'Table 7: Calibration Ablation Study'),
        ('Table8_Action_Grids.csv', 'Table 8: Alternative Action Grids'),
        ('Table9_Alpha_Sensitivity.csv', 'Table 9: Exploration Parameter (α) Sensitivity'),
        ('Table10_Nonlinear_Policies.csv', 'Table 10: Nonlinear Contextual Policies'),
        ('Table11_Random_Context.csv', 'Table 11: Randomized-Context Placebo Test'),
        ('Table12_Shuffled_Action.csv', 'Table 12: Shuffled-Action Placebo Test'),
        ('Table13_Temporal_Subsamples.csv', 'Table 13: Temporal Subsample Analysis'),
        ('Table14_Source_vs_Target.csv', 'Table 14: Source vs. Target Degeneracy'),
        ('Table15_Robustness_Summary.csv', 'Table 15: Robustness Summary'),
    ]
    
    for filename, caption in table_configs:
        filepath = os.path.join(data_dir, filename)
        if os.path.exists(filepath):
            df = pd.read_csv(filepath)
            add_dataframe_to_doc(doc, df, caption)
            print(f"  ✅ 添加 {caption}")
        else:
            print(f"  ❌ 未找到 {filename}")
    
    # 保存文档
    output_file = "/Volumes/agents/codex/paper2-LSTM量化金融/archive_20260710_1046_chatgpt_review/JFDS_Submission_Package_20260712/paper/All_Tables_Complete.docx"
    doc.save(output_file)
    
    print(f"\n✅ 所有表格已保存到: {output_file}")

if __name__ == '__main__':
    main()
