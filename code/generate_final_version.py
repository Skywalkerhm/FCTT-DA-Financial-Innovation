import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# Load the original document
doc = docx.Document("main_paper_jfds_ready.docx")

# Function to add expanded content from markdown
def add_markdown_content(doc, markdown_text, start_after_heading=None):
    """Add markdown content to the document after a specific heading."""
    lines = markdown_text.strip().split('\n')
    
    # Find the position to insert
    insert_index = None
    if start_after_heading:
        for i, para in enumerate(doc.paragraphs):
            if start_after_heading in para.text:
                insert_index = i + 1
                break
    
    if insert_index is None:
        print(f"Warning: Could not find heading '{start_after_heading}'")
        return
    
    # Remove existing content until next major heading
    end_index = insert_index
    for i in range(insert_index, len(doc.paragraphs)):
        text = doc.paragraphs[i].text.strip()
        style = doc.paragraphs[i].style.name if doc.paragraphs[i].style else ""
        
        # Stop at next major section (numbered heading)
        if re.match(r'^\d+\.\s+', text) and 'Heading 2' in style:
            end_index = i
            break
        # Also stop at References or Appendix
        if text in ['References'] or text.startswith('Appendix'):
            end_index = i
            break
        end_index = i + 1
    
    # Delete existing paragraphs in range
    for _ in range(end_index - insert_index):
        p = doc.paragraphs[insert_index]
        p._element.getparent().remove(p._element)
    
    # Add new content
    current_pos = insert_index
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Determine style based on markdown formatting
        if line.startswith('## '):
            # Main heading
            p = doc.add_paragraph(line[3:], style='Heading 2')
            p._element.getparent().insert(current_pos, p._element)
            current_pos += 1
        elif line.startswith('### '):
            # Sub heading
            p = doc.add_paragraph(line[4:], style='Heading 3')
            p._element.getparent().insert(current_pos, p._element)
            current_pos += 1
        elif line.startswith('**') and line.endswith('**'):
            # Bold text (likely a sub-sub heading)
            p = doc.add_paragraph(line[2:-2], style='Normal')
            p.runs[0].bold = True
            p._element.getparent().insert(current_pos, p._element)
            current_pos += 1
        elif line.startswith('- '):
            # Bullet point
            p = doc.add_paragraph(line[2:], style='List Bullet')
            p._element.getparent().insert(current_pos, p._element)
            current_pos += 1
        elif re.match(r'^\d+\.\s+', line):
            # Numbered list
            p = doc.add_paragraph(line, style='List Number')
            p._element.getparent().insert(current_pos, p._element)
            current_pos += 1
        elif line.startswith('|') and '|' in line[1:]:
            # Table row - skip for now, handle tables separately
            continue
        elif line.startswith('**Table'):
            # Table caption
            p = doc.add_paragraph(line, style='Normal')
            p.runs[0].bold = True
            p._element.getparent().insert(current_pos, p._element)
            current_pos += 1
        else:
            # Regular paragraph
            p = doc.add_paragraph(line, style='Normal')
            p._element.getparent().insert(current_pos, p._element)
            current_pos += 1

# Read expanded content files
with open("related_work_expanded.md", "r") as f:
    related_work_content = f.read()

with open("methodology_expanded.md", "r") as f:
    methodology_content = f.read()

with open("results_expanded.md", "r") as f:
    results_content = f.read()

print("Loaded expanded content files")

# Note: This is a simplified version - in practice, you'd need to handle
# the document structure more carefully
print("Note: Manual integration of expanded content is recommended")
print("Expanded content files are ready for integration:")
print("  - related_work_expanded.md (1951 words)")
print("  - methodology_expanded.md (1994 words)")
print("  - results_expanded.md (1813 words)")

