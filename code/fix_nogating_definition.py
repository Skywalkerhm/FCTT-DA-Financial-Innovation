import docx
from docx.shared import Pt
import re

# 加载文档
doc = docx.Document("/Volumes/agents/codex/paper2-LSTM量化金融/archive_20260710_1046_chatgpt_review/JFDS_Submission_Package_20260712/paper/论文_公式恢复与图表美化_TABLE4_FIXED.docx")

print("=== 检查 NoGating 定义冲突 ===\n")

# 1. 查找所有提到 NoGating 的地方
print("1. 查找所有提到 NoGating 的地方:")
nogating_refs = []
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if 'NoGating' in text or 'no gating' in text.lower() or 'no-gating' in text.lower():
        nogating_refs.append((i, text[:150]))
        print(f"   段落 {i}: {text[:150]}...")

print(f"\n   总共找到 {len(nogating_refs)} 处")

# 2. 查找 τ=1.00 的定义
print("\n2. 查找 τ=1.00 的定义:")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '1.00' in text and ('never' in text.lower() or 'no-trade' in text.lower() or 'no trade' in text.lower()):
        print(f"   段落 {i}: {text[:150]}...")

# 3. 查找 τ=0.50 的定义
print("\n3. 查找 τ=0.50 的定义:")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '0.50' in text and ('static' in text.lower() or 'threshold' in text.lower()):
        print(f"   段落 {i}: {text[:150]}...")

# 4. 查找 Table 3、5、6 中的 NoGating
print("\n4. 查找表格中的 NoGating:")
for i, table in enumerate(doc.tables):
    for row in table.rows:
        for cell in row.cells:
            if 'NoGating' in cell.text:
                print(f"   Table {i+1}: {cell.text[:100]}...")

