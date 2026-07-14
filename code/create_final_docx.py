import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def read_markdown_file(filepath):
    """Read a markdown file and return its content."""
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()

def parse_markdown_sections(content):
    """Parse markdown content into sections."""
    sections = []
    current_section = None
    current_content = []
    
    for line in content.split('\n'):
        if line.startswith('## '):
            if current_section:
                sections.append({
                    'title': current_section,
                    'content': '\n'.join(current_content)
                })
            current_section = line[3:].strip()
            current_content = []
        elif line.startswith('### '):
            current_content.append(line)
        elif line.strip():
            current_content.append(line)
    
    if current_section:
        sections.append({
            'title': current_section,
            'content': '\n'.join(current_content)
        })
    
    return sections

def add_section_to_doc(doc, section_content, heading_level=2):
    """Add a section's content to the document."""
    lines = section_content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('**Table'):
            p = doc.add_paragraph(line)
            if p.runs:
                p.runs[0].bold = True
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.\s+', line):
            doc.add_paragraph(line, style='List Number')
        elif line.startswith('|'):
            # Skip table rows for now
            continue
        else:
            doc.add_paragraph(line)

# Load original document
original_doc = docx.Document("main_paper_jfds_ready.docx")

# Create new document
new_doc = docx.Document()

# Set default font
style = new_doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Set margins
for section in new_doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Copy title and author info
for para in original_doc.paragraphs[:10]:
    text = para.text.strip()
    if text:
        if para.style and 'Heading' in para.style.name:
            new_doc.add_heading(text, level=1)
        else:
            new_doc.add_paragraph(text)

# Add abstract
new_doc.add_heading('Abstract', level=1)
abstract_text = """We identify Static-Rule Collapse, a silent failure mode in cross-asset decision-layer transfer where a frozen contextual policy degenerates into a fixed probability threshold on out-of-distribution targets. This failure mode is undetectable by standard backtests, as the collapsed policy may still outperform naive baselines while losing all contextual adaptation capability.

Across 234 source-target pairs (6 sources × 39 targets), we establish an empirical taxonomy: degenerate sources (SPY, QQQ, EEM, GLD) trigger 100% target collapse; diverse sources (TLT, IWM) maintain contextual adaptation. We prove that source diversity is a necessary condition for transfer and provide sufficient conditions linking source-side training dynamics to target-side failure.

We contribute FCTT-DA (Frozen Contextual Threshold Transfer with Degeneracy Audit), a deployable framework with source diversity pre-checks, behavioral-equivalent baseline comparison, transaction cost sensitivity analysis, and an action-region geometry diagnostic. The framework detects collapse before deployment, preventing futile transfers and saving computational resources.

Our findings have immediate implications for quantitative trading firms deploying machine learning policies across multiple assets. We demonstrate that the common practice of transferring policies from single assets is fundamentally flawed, and propose source diversity as a prerequisite for successful cross-asset transfer."""
new_doc.add_paragraph(abstract_text)

# Add keywords (adjusted to 5)
new_doc.add_paragraph('Keywords: contextual bandits, policy transfer, decision gating, static-rule collapse, cross-asset learning')

# Add Introduction
new_doc.add_heading('1. Introduction', level=1)

# Copy introduction content from original
in_intro = False
for para in original_doc.paragraphs:
    text = para.text.strip()
    style_name = para.style.name if para.style else ""
    
    if "1. Introduction" in text and "Heading" in style_name:
        in_intro = True
        continue
    
    if in_intro and "2. Related Work" in text and "Heading" in style_name:
        break
    
    if in_intro and text:
        if "Heading" in style_name:
            new_doc.add_heading(text, level=2)
        else:
            new_doc.add_paragraph(text)

# Add expanded Related Work
related_work_content = read_markdown_file("related_work_expanded.md")
related_work_sections = parse_markdown_sections(related_work_content)

new_doc.add_heading('2. Related Work', level=1)
for section in related_work_sections:
    if section['title'] != '2. Related Work':
        new_doc.add_heading(section['title'], level=2)
    add_section_to_doc(new_doc, section['content'])

# Add expanded Methodology
methodology_content = read_markdown_file("methodology_expanded.md")
methodology_sections = parse_markdown_sections(methodology_content)

for section in methodology_sections:
    new_doc.add_heading(section['title'], level=1)
    add_section_to_doc(new_doc, section['content'])

# Add Experimental Design (from original)
new_doc.add_heading('6. Experimental Design', level=1)
in_experimental = False
for para in original_doc.paragraphs:
    text = para.text.strip()
    style_name = para.style.name if para.style else ""
    
    if "6. Experimental Design" in text and "Heading" in style_name:
        in_experimental = True
        continue
    
    if in_experimental and "7. Results" in text and "Heading" in style_name:
        break
    
    if in_experimental and text:
        if "Heading" in style_name:
            new_doc.add_heading(text, level=2)
        else:
            new_doc.add_paragraph(text)

# Add expanded Results
results_content = read_markdown_file("results_expanded.md")
results_sections = parse_markdown_sections(results_content)

new_doc.add_heading('7. Results', level=1)
for section in results_sections:
    if section['title'] != '7. Results':
        new_doc.add_heading(section['title'], level=2)
    add_section_to_doc(new_doc, section['content'])

# Add Discussion and Conclusion (from original)
new_doc.add_heading('8. Discussion', level=1)
in_discussion = False
for para in original_doc.paragraphs:
    text = para.text.strip()
    style_name = para.style.name if para.style else ""
    
    if "8. Discussion" in text and "Heading" in style_name:
        in_discussion = True
        continue
    
    if in_discussion and "9. Conclusion" in text and "Heading" in style_name:
        break
    
    if in_discussion and text:
        if "Heading" in style_name:
            new_doc.add_heading(text, level=2)
        else:
            new_doc.add_paragraph(text)

# Add Conclusion
new_doc.add_heading('9. Conclusion', level=1)
in_conclusion = False
for para in original_doc.paragraphs:
    text = para.text.strip()
    style_name = para.style.name if para.style else ""
    
    if "9. Conclusion" in text and "Heading" in style_name:
        in_conclusion = True
        continue
    
    if in_conclusion and "JEL Classification" in text:
        break
    
    if in_conclusion and text:
        if "Heading" in style_name:
            new_doc.add_heading(text, level=2)
        else:
            new_doc.add_paragraph(text)

# Add JEL Classification and other metadata
new_doc.add_heading('JEL Classification', level=2)
new_doc.add_paragraph('G11 (Portfolio Choice); G12 (Asset Pricing); C44 (Statistical Decision Theory); C61 (Optimization Techniques)')

new_doc.add_heading('Data Availability Statement', level=2)
new_doc.add_paragraph('All data sources are publicly available. Equity and ETF price data are from Yahoo Finance. Cryptocurrency data are from public blockchain explorers. Complete data manifests with SHA-256 checksums are provided in the submission package.')

new_doc.add_heading('Conflict of Interest', level=2)
new_doc.add_paragraph('The author declares no conflicts of interest.')

# Add Acknowledgments (NEW)
new_doc.add_heading('Acknowledgments', level=1)
new_doc.add_paragraph('The author感谢编辑和匿名审稿人的宝贵建议。')

# Add Author Contributions (NEW)
new_doc.add_heading('Author Contributions', level=1)
new_doc.add_paragraph('Min Huang: Conceptualization, Methodology, Software, Validation, Formal analysis, Investigation, Data curation, Writing - original draft, Writing - review & editing, Visualization, Supervision, Project management.')

# Add References (from original, with formatting fixes)
new_doc.add_heading('References', level=1)
in_references = False
for para in original_doc.paragraphs:
    text = para.text.strip()
    if "References" in text and len(text) < 20:
        in_references = True
        continue
    
    if in_references and text:
        # Fix reference formatting
        ref_text = text
        # Remove leading dash
        if ref_text.startswith('- '):
            ref_text = ref_text[2:]
        # Replace "and" with "&"
        ref_text = ref_text.replace(', and ', ' & ')
        new_doc.add_paragraph(ref_text)

# Save the document
output_filename = "main_paper_jfds_READY_v2.docx"
new_doc.save(output_filename)
print(f"Created final document: {output_filename}")

# Count words
total_words = 0
for para in new_doc.paragraphs:
    words = len(re.findall(r'\w+', para.text))
    total_words += words

print(f"Total words: {total_words}")
print(f"Estimated main content words: {total_words - 1500}")  # Subtract references and appendices

