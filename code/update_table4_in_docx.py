import docx
import pandas as pd

# 加载文档
doc = docx.Document("/Volumes/agents/codex/paper2-LSTM量化金融/archive_20260710_1046_chatgpt_review/JFDS_Submission_Package_20260712/paper/论文_公式恢复与图表美化_FINAL.docx")

print("=== 更新 Table 4 ===\n")

# 读取修正后的数据
corrected_data = pd.read_csv("/Volumes/agents/codex/paper2-LSTM量化金融/archive_20260710_1046_chatgpt_review/JFDS_Submission_Package_20260712/paper/table_data/Table4_Collapse_vs_Threshold_CORRECTED.csv")

print("修正后的数据:")
print(corrected_data.to_string(index=False))

# 找到Table 4并更新
print("\n查找Table 4...")
table4_found = False

for i, table in enumerate(doc.tables):
    # 检查是否是Table 4
    header_row = table.rows[0]
    header_text = [cell.text.strip() for cell in header_row.cells]
    
    if 'Modal Share Threshold' in header_text or 'Modal_share_threshold' in header_text:
        print(f"找到Table {i+1}: {header_text}")
        table4_found = True
        
        # 更新数据行
        for row_idx, (_, data_row) in enumerate(corrected_data.iterrows()):
            if row_idx < len(table.rows) - 1:  # 跳过表头
                row = table.rows[row_idx + 1]
                
                # 更新每个单元格
                for col_idx, cell in enumerate(row.cells):
                    col_name = corrected_data.columns[col_idx]
                    new_value = str(data_row[col_name])
                    old_value = cell.text.strip()
                    
                    if old_value != new_value:
                        # 清除并更新
                        for para in cell.paragraphs:
                            para.clear()
                            para.add_run(new_value)
                        print(f"  更新 Row {row_idx+1}, Col {col_idx}: {old_value} -> {new_value}")
        
        break

if not table4_found:
    print("未找到Table 4，尝试其他方法...")
    
    # 查找包含collapse rate数据的表格
    for i, table in enumerate(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                if '79%' in cell.text or '185' in cell.text:
                    print(f"找到包含79%/185的表格: Table {i+1}")
                    print(f"  单元格内容: {cell.text[:50]}...")

# 保存文档
output_file = "/Volumes/agents/codex/paper2-LSTM量化金融/archive_20260710_1046_chatgpt_review/JFDS_Submission_Package_20260712/paper/论文_公式恢复与图表美化_TABLE4_FIXED.docx"
doc.save(output_file)

print(f"\n✅ Table 4 更新完成！")
print(f"保存为: {output_file}")

