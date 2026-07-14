import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os
import re

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color
    })
    shading_elm.append(shading)

def add_table_after_paragraph(doc, para, df, font_size=8):
    """在指定段落后添加表格"""
    # 创建表格
    rows, cols = df.shape
    table = doc.add_table(rows=rows+1, cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 添加表头
    header = table.rows[0]
    for i, col_name in enumerate(df.columns):
        cell = header.cells[i]
        cell.text = str(col_name)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(font_size)
                run.font.name = 'Times New Roman'
        set_cell_shading(cell, 'D9E2F3')
    
    # 添加数据行
    for i, (_, row) in enumerate(df.iterrows()):
        for j, value in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
                    run.font.name = 'Times New Roman'
    
    # 将表格移到段落之后
    para._element.addnext(table._element)
    
    return table

def main():
    # 加载表格数据
    data_dir = "/Volumes/agents/codex/paper2-LSTM量化金融/archive_20260710_1046_chatgpt_review/JFDS_Submission_Package_20260712/paper/table_data"
    
    # 加载主文档
    main_doc_path = "/Volumes/agents/codex/paper2-LSTM量化金融/archive_20260710_1046_chatgpt_review/JFDS_Submission_Package_20260712/paper/main_paper_jfds_READY_v20_formatted.docx"
    doc = Document(main_doc_path)
    
    # 定义表格映射（表名 -> 数据文件）
    table_mapping = {
        'Table 1': 'Table1_Source_Diversity.csv',
        'Table 2': 'Table2_Collapse_by_Source.csv',
        'Table 3': 'Table3_TLT_Performance.csv',
        'Table 4': 'Table4_Collapse_vs_Threshold.csv',
        'Table 5': 'Table5_Performance_by_Class.csv',
        'Table 6': 'Table6_Transaction_Costs.csv',
        'Table 7': 'Table7_Calibration_Ablation.csv',
        'Table 8': 'Table8_Action_Grids.csv',
        'Table 9': 'Table9_Alpha_Sensitivity.csv',
        'Table 10': 'Table10_Nonlinear_Policies.csv',
        'Table 11': 'Table11_Random_Context.csv',
        'Table 12': 'Table12_Shuffled_Action.csv',
        'Table 13': 'Table13_Temporal_Subsamples.csv',
        'Table 14': 'Table14_Source_vs_Target.csv',
        'Table 15': 'Table15_Robustness_Summary.csv',
    }
    
    # 查找并替换表格
    tables_added = 0
    paragraphs_to_process = []
    
    # 先收集需要处理的段落
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        for table_name, data_file in table_mapping.items():
            if text.startswith(table_name) and ':' in text:
                paragraphs_to_process.append((i, table_name, data_file))
    
    # 从后往前处理，避免索引问题
    for i, table_name, data_file in reversed(paragraphs_to_process):
        filepath = os.path.join(data_dir, data_file)
        if os.path.exists(filepath):
            df = pd.read_csv(filepath)
            para = doc.paragraphs[i]
            add_table_after_paragraph(doc, para, df)
            tables_added += 1
            print(f"  ✅ 添加 {table_name}")
    
    # 保存文档
    output_file = "/Volumes/agents/codex/paper2-LSTM量化金融/archive_20260710_1046_chatgpt_review/JFDS_Submission_Package_20260712/paper/main_paper_jfds_READY_v21_with_tables.docx"
    doc.save(output_file)
    
    print(f"\n✅ 已添加 {tables_added} 个表格到主文档")
    print(f"保存为: {output_file}")

if __name__ == '__main__':
    main()
