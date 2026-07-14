import docx
import re

def fix_text(text, old_patterns, new_patterns):
    """Replace old patterns with new patterns in text."""
    for old, new in zip(old_patterns, new_patterns):
        text = text.replace(old, new)
    return text

# Load the document
doc = docx.Document("main_paper_jfds_READY_v12_complete.docx")

print("=== 修复数字不一致 ===\n")

# 定义正确的数字
# 6个源资产，39个目标资产，234对（6×39）
# 资产类别：15 equity + 7 fixed income + 6 commodity + 3 real estate + 6 FX + 2 crypto = 39

# 需要替换的模式
old_patterns = [
    '156 source-target pairs',
    '117 pairs',
    'Targets (39): 15 equity, 7 fixed income, 6 commodity, 3 real estate, 6 FX, 3 crypto',
    'three of the four source policies',
    'three of four sources',
]

new_patterns = [
    '234 source-target pairs',
    '234 pairs',
    'Targets (39): 15 equity, 7 fixed income, 6 commodity, 3 real estate, 6 FX, 2 crypto',
    'three of the six source policies',
    'three of six sources',
]

# 修复文档
fixed_count = 0
for i, para in enumerate(doc.paragraphs):
    original_text = para.text
    fixed_text = original_text
    
    # 修复数字
    for old, new in zip(old_patterns, new_patterns):
        if old in fixed_text:
            fixed_text = fixed_text.replace(old, new)
            fixed_count += 1
    
    # 修复四源资产到六源资产
    if 'four source assets' in fixed_text and 'SPY, QQQ, EEM' in fixed_text:
        fixed_text = fixed_text.replace('four source assets', 'six source assets')
        fixed_count += 1
    
    if 'four source' in fixed_text and 'degenerate' in fixed_text:
        fixed_text = fixed_text.replace('four source', 'six source')
        fixed_count += 1
    
    # 更新段落
    if fixed_text != original_text:
        para.clear()
        para.add_run(fixed_text)
        print(f"  修复段落 {i}: {original_text[:60]}... -> {fixed_text[:60]}...")

# 保存文档
output_file = "main_paper_jfds_READY_v13_numbers_fixed.docx"
doc.save(output_file)

print(f"\n✅ 数字修复完成！")
print(f"修复数量: {fixed_count}")
print(f"保存为: {output_file}")

