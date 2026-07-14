import docx
from docx.shared import Pt
import re

# 加载文档
doc = docx.Document("/Volumes/agents/codex/paper2-LSTM量化金融/archive_20260710_1046_chatgpt_review/JFDS_Submission_Package_20260712/paper/论文_公式恢复与图表美化_TABLE4_FIXED.docx")

print("=== 修正 NoGating 定义冲突 ===\n")

# 定义修正规则
# 核心问题：
# - τ=0.50 应该称为 "Ungated classifier" 或 "Static-0.50"
# - τ=1.00 应该称为 "No-trade baseline"
# - "NoGating" 这个术语应该统一为 τ=0.50

replacements = [
    # 修正 τ=1.00 的错误定义
    ('threshold 1.00 corresponds to "NoGating"', 'threshold 1.00 corresponds to "No-trade baseline"'),
    ('threshold 1.00 corresponds to NoGating', 'threshold 1.00 corresponds to No-trade baseline'),
    ('"NoGating" (never taking a position)', '"No-trade baseline" (never taking a position)'),
    ('NoGating (never taking a position)', 'No-trade baseline (never taking a position)'),
    
    # 统一 NoGating 为 τ=0.50
    ('NoGating: A static threshold of', 'NoGating: A static threshold of 0.50,'),
    ('NoGating baseline (static threshold)', 'NoGating baseline (static 0.50 threshold)'),
    ('NoGating baseline (static 0.50 threshold)', 'NoGating baseline (static 0.50 threshold)'),
    
    # 修正表格中的描述
    ('ZS6 vs. NoGating (0.50)', 'ZS6 vs. Static-0.50'),
    ('Mean Benefit vs. NoGating', 'Mean Benefit vs. Static-0.50'),
    ('Mean Sharpe vs. NoGating', 'Mean Sharpe vs. Static-0.50'),
    
    # 添加定义说明
    ('NoGating: A static threshold of 0.50,', 'NoGating: A static threshold of 0.50 (also called "Ungated classifier"),'),
]

# 修正文档
fixed_count = 0
for i, para in enumerate(doc.paragraphs):
    original_text = para.text
    fixed_text = original_text
    
    # 应用替换
    for old, new in replacements:
        if old in fixed_text:
            fixed_text = fixed_text.replace(old, new)
            fixed_count += 1
    
    # 更新段落
    if fixed_text != original_text:
        para.clear()
        para.add_run(fixed_text)
        print(f"  修正段落 {i}:")
        print(f"    原文: {original_text[:80]}...")
        print(f"    修正: {fixed_text[:80]}...")

# 修正表格
print("\n修正表格:")
for i, table in enumerate(doc.tables):
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            original_text = cell.text
            fixed_text = original_text
            
            # 应用替换
            for old, new in replacements:
                if old in fixed_text:
                    fixed_text = fixed_text.replace(old, new)
                    fixed_count += 1
            
            # 更新单元格
            if fixed_text != original_text:
                for para in cell.paragraphs:
                    para.clear()
                    para.add_run(fixed_text)
                print(f"  修正 Table {i+1}, Row {row_idx}, Cell {cell_idx}:")
                print(f"    原文: {original_text[:80]}...")
                print(f"    修正: {fixed_text[:80]}...")

# 保存文档
output_file = "/Volumes/agents/codex/paper2-LSTM量化金融/archive_20260710_1046_chatgpt_review/JFDS_Submission_Package_20260712/paper/论文_公式恢复与图表美化_NOGATING_FIXED.docx"
doc.save(output_file)

print(f"\n✅ NoGating 定义修正完成！")
print(f"修正数量: {fixed_count}")
print(f"保存为: {output_file}")

