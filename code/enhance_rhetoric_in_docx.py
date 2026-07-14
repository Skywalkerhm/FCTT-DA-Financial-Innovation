import docx
from docx.shared import Pt
import re

# 加载文档
doc = docx.Document("/Volumes/agents/codex/paper2-LSTM量化金融/archive_20260710_1046_chatgpt_review/JFDS_Submission_Package_20260712/paper/论文_公式恢复与图表美化_FINAL.docx")

print("=== 增强修辞效果 ===\n")

# 1. 修改Abstract
print("1. 修改Abstract...")

# 找到Abstract段落并替换
abstract_found = False
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text == 'Abstract':
        abstract_found = True
        # 找到Abstract后的段落
        for j in range(i+1, min(i+10, len(doc.paragraphs))):
            next_para = doc.paragraphs[j]
            if next_para.text.strip():
                # 替换为增强版
                new_abstract = """We identify Source Static Degeneracy, a silent failure mode in cross-asset decision-layer transfer where a frozen contextual policy degenerates into a fixed probability threshold. This degeneracy remains invisible to standard backtesting protocols: the collapsed policy may still outperform a naive static threshold, masking the complete loss of contextual adaptation. A quantitative team deploying such a policy would conclude that "transfer works," unaware that the policy has abandoned context-dependent decision-making entirely.

Across 234 source-target pairs spanning six asset classes, we establish an empirical taxonomy: degenerate sources (SPY, QQQ, EEM, GLD) trigger universal action concentration on 100% of targets; diverse sources (TLT, IWM) maintain behavioral variation. We prove that source-side action concentration is associated with target-side degeneracy and provide sufficient conditions linking source diversity to transfer outcomes.

We contribute FCTT-DA (Frozen Contextual Threshold Transfer with Degeneracy Audit), a deployable framework with source diversity pre-checks, behavioral-equivalent baseline comparison, and transaction cost sensitivity analysis. The framework is designed to detect what backtests cannot: policies that appear to adapt but are behaviorally static.

Our findings challenge the implicit assumption that contextual policies transfer contextual behavior. The question is not whether a policy transfers, but whether it transfers adaptation. FCTT-DA provides the diagnostic tools to answer this question before deployment."""
                
                # 清除并替换
                next_para.clear()
                next_para.add_run(new_abstract)
                print(f"   ✅ 替换Abstract (段落 {j})")
                break
        break

# 2. 在Introduction末尾添加新段落
print("\n2. 添加 'The Narrative Task' 段落...")

# 找到1.3 Contributions末尾
in_contributions = False
contributions_end = None

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '1.3' in text and 'Contributions' in text:
        in_contributions = True
    if in_contributions and ('2.' in text or 'Related Work' in text):
        contributions_end = i
        break

if contributions_end:
    # 在Contributions末尾插入新段落
    new_section = """1.4 The Narrative Task: From Optimism to Audit

The financial machine learning literature has developed an implicit optimism about cross-asset transfer. The prevailing narrative assumes that if a prediction model transfers well across assets, the decision layer will follow. This paper challenges that assumption.

We propose a reframing of the research question:

• Prevailing: "Can this policy transfer to asset X?" → Ours: "Does this policy retain contextual adaptation on asset X?"
• Prevailing: "Does the transferred policy outperform 0.50?" → Ours: "Does the transferred policy outperform its behavioral equivalent?"
• Prevailing: "Is the Sharpe ratio positive?" → Ours: "Is the action distribution diverse?"
• Prevailing: "Does transfer work?" → Ours: "Is decision diversity sufficient?"

This reframing shifts the focus from outcome evaluation to behavioral audit. A policy that outperforms a naive baseline may still be behaviorally static—and a static policy, however profitable, has abandoned the very adaptation that justified using a contextual bandit in the first place.

The practical implication is profound: standard backtesting protocols cannot distinguish between genuine contextual transfer and inherited static degeneracy. This creates a blind spot in model validation that affects any firm deploying machine learning policies across multiple assets.

FCTT-DA is designed to fill this blind spot. By auditing source-side diversity before transfer and verifying behavioral equivalence after transfer, the framework provides what backtests cannot: a diagnostic for decision-layer adaptation."""
    
    # 在指定位置插入
    para = doc.paragraphs[contributions_end]
    new_para = doc.add_paragraph()
    new_para.style = doc.styles['Heading 2']
    new_para.add_run(new_section)
    
    # 在当前位置前插入
    para._element.addprevious(new_para._element)
    print(f"   ✅ 添加 'The Narrative Task' 段落")

# 3. 增强关键句子
print("\n3. 增强关键句子...")

enhancements = {
    'We identify Source Static Degeneracy, a failure mode': 
        'We identify Source Static Degeneracy, a silent failure mode',
    'a failure mode in cross-asset decision-layer transfer':
        'a silent failure mode in cross-asset decision-layer transfer where a frozen contextual policy degenerates into a fixed probability threshold',
}

for i, para in enumerate(doc.paragraphs):
    original_text = para.text
    for old, new in enhancements.items():
        if old in original_text:
            new_text = original_text.replace(old, new)
            para.clear()
            para.add_run(new_text)
            print(f"   ✅ 增强段落 {i}")
            break

# 保存文档
output_file = "/Volumes/agents/codex/paper2-LSTM量化金融/archive_20260710_1046_chatgpt_review/JFDS_Submission_Package_20260712/paper/论文_公式恢复与图表美化_RHETORIC.docx"
doc.save(output_file)

print(f"\n✅ 修辞效果增强完成！")
print(f"保存为: {output_file}")

