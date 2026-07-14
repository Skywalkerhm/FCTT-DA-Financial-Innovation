import docx
from docx.shared import Pt
import re

def read_markdown_file(filepath):
    """Read a markdown file and return its content."""
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()

def add_ablation_to_doc(doc, markdown_content):
    """Add ablation studies to the document."""
    lines = markdown_content.strip().split('\n')
    
    # Find the position to insert (after 7.7 or before 8. Discussion)
    insert_index = None
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if '8. Discussion' in text:
            insert_index = i
            break
    
    if insert_index is None:
        print("❌ 未找到插入位置")
        return False
    
    # Add content before Discussion section
    current_pos = insert_index
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        new_para = doc.add_paragraph()
        
        # Determine style based on markdown formatting
        if line.startswith('## '):
            new_para.style = doc.styles['Heading 1']
            new_para.add_run(line[3:])
        elif line.startswith('### '):
            new_para.style = doc.styles['Heading 2']
            new_para.add_run(line[4:])
        elif line.startswith('**') and line.endswith('**'):
            run = new_para.add_run(line[2:-2])
            run.bold = True
        elif line.startswith('**Table'):
            run = new_para.add_run(line)
            run.bold = True
        elif line.startswith('- '):
            new_para.style = doc.styles['List Bullet']
            new_para.add_run(line[2:])
        elif re.match(r'^\d+\.\s+', line):
            new_para.style = doc.styles['List Number']
            new_para.add_run(line)
        elif line.startswith('|'):
            # Table row - skip for now
            continue
        else:
            new_para.add_run(line)
        
        # Insert before Discussion
        doc.paragraphs[insert_index]._element.addprevious(new_para._element)
        insert_index += 1
    
    return True

# Load the document
doc = docx.Document("main_paper_jfds_READY_v18_claims_adjusted.docx")

print("=== 添加消融实验章节 ===\n")

# Read the ablation studies content
ablation_content = read_markdown_file("ablation_studies.md")

# Add to document
success = add_ablation_to_doc(doc, ablation_content)

if success:
    # Save the document
    output_file = "main_paper_jfds_READY_v19_ablation.docx"
    doc.save(output_file)
    
    print(f"✅ 消融实验章节添加完成！")
    print(f"保存为: {output_file}")
    
    # Count total words
    total_words = sum(len(re.findall(r'\w+', para.text)) for para in doc.paragraphs)
    print(f"文档总字数: {total_words}")
else:
    print("❌ 添加失败")

