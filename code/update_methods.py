import docx
from docx.shared import Pt
import re

def read_markdown_file(filepath):
    """Read a markdown file and return its content."""
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()

def add_content_to_section(doc, section_keyword, new_content, after=True):
    """Add content to a specific section."""
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if section_keyword in text:
            # Found the section
            lines = new_content.strip().split('\n')
            
            # Add content after the section header
            last_para = para
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                new_para = doc.add_paragraph()
                
                if line.startswith('### '):
                    new_para.style = doc.styles['Heading 3']
                    new_para.add_run(line[4:])
                elif line.startswith('**') and line.endswith('**'):
                    run = new_para.add_run(line[2:-2])
                    run.bold = True
                elif line.startswith('- '):
                    new_para.style = doc.styles['List Bullet']
                    new_para.add_run(line[2:])
                elif re.match(r'^\d+\.\s+', line):
                    new_para.style = doc.styles['List Number']
                    new_para.add_run(line)
                else:
                    new_para.add_run(line)
                
                last_para._element.addnext(new_para._element)
                last_para = new_para
            
            print(f"  ✅ 添加内容到 {section_keyword}")
            return True
    
    print(f"  ❌ 未找到 {section_keyword}")
    return False

# Load the document
doc = docx.Document("main_paper_jfds_READY_v16_cleaned.docx")

print("=== 更新方法描述 ===\n")

# Read the supplement
supplement = read_markdown_file("method_details_supplement.md")

# Parse the supplement into sections
sections = {}
current_section = None
current_content = []

for line in supplement.split('\n'):
    if line.startswith('## '):
        if current_section:
            sections[current_section] = '\n'.join(current_content)
        current_section = line[3:].strip()
        current_content = []
    else:
        current_content.append(line)

if current_section:
    sections[current_section] = '\n'.join(current_content)

# Add content to relevant sections
print("1. 更新6.1 Data部分...")
if '6.1 Data (Expanded)' in sections:
    add_content_to_section(doc, '6.1 Data', sections['6.1 Data (Expanded)'])

print("\n2. 更新6.2 Features部分...")
if '6.2 Features (Expanded)' in sections:
    add_content_to_section(doc, '6.2 Features', sections['6.2 Features (Expanded)'])

print("\n3. 更新6.3 Protocol部分...")
if '6.3 Protocol (Expanded)' in sections:
    add_content_to_section(doc, '6.3 Protocol', sections['6.3 Protocol (Expanded)'])

print("\n4. 更新6.4 Execution部分...")
if '6.4 Execution (Expanded)' in sections:
    add_content_to_section(doc, '6.4 Execution', sections['6.4 Execution (Expanded)'])

print("\n5. 更新6.5 Reproducibility部分...")
if '6.5 Reproducibility (Expanded)' in sections:
    add_content_to_section(doc, '6.5 Reproducibility', sections['6.5 Reproducibility (Expanded)'])

# Save the document
output_file = "main_paper_jfds_READY_v17_methods_updated.docx"
doc.save(output_file)

print(f"\n✅ 方法描述更新完成！")
print(f"保存为: {output_file}")

