import docx
import re

def fix_text(text, old_patterns, new_patterns):
    """替换文本中的模式"""
    for old, new in zip(old_patterns, new_patterns):
        text = text.replace(old, new)
    return text

# 加载文档
doc = docx.Document("/Volumes/agents/codex/paper2-LSTM量化金融/archive_20260710_1046_chatgpt_review/JFDS_Submission_Package_20260712/paper/main_paper_jfds_READY_v20_formatted.docx")

print("=== 修正文档中的数字 ===\n")

# 需要修正的模式
old_patterns = [
    # 旧的资产分类
    '15 equity, 7 fixed income, 6 commodity, 3 real estate, 6 FX, 2 crypto',
    '15 equity + 7 fixed income + 6 commodity + 3 real estate + 6 FX + 2 crypto = 39',
    '15 equity + 7 fixed income + 6 commodity + 3 real estate + 6 FX + 3 crypto',
    '15+7+6+3+6+2=39',
    '15+7+6+3+6+3=40',
    
    # 不一致的数字
    '156 source-target pairs',
    '117 pairs',
    '100% of 117 pairs',
    
    # 错误的资产分类
    'US Equity ETFs (12)',
    'International Equity ETFs (8)',
    'Fixed Income ETFs (6)',
    'Commodity ETFs (5)',
    'Real Estate ETFs (3)',
    'Currency ETFs (3)',
    'Crypto Assets (3)',
]

new_patterns = [
    # 新的资产分类
    '11 US equity, 7 international equity, 6 fixed income, 5 commodity, 3 real estate, 3 currency, 4 crypto',
    '11 US equity + 7 international equity + 6 fixed income + 5 commodity + 3 real estate + 3 currency + 4 crypto = 39',
    '11 US equity + 7 international equity + 6 fixed income + 5 commodity + 3 real estate + 3 currency + 4 crypto',
    '11+7+6+5+3+3+4=39',
    '11+7+6+5+3+3+4=39',
    
    # 统一的数字
    '234 source-target pairs',
    '234 pairs',
    '100% of 234 pairs',
    
    # 正确的资产分类
    'US Equity ETFs (11)',
    'International Equity ETFs (7)',
    'Fixed Income ETFs (6)',
    'Commodity ETFs (5)',
    'Real Estate ETFs (3)',
    'Currency ETFs (3)',
    'Crypto Assets (4)',
]

# 修正文档
fixed_count = 0
for i, para in enumerate(doc.paragraphs):
    original_text = para.text
    fixed_text = original_text
    
    # 替换模式
    for old, new in zip(old_patterns, new_patterns):
        if old in fixed_text:
            fixed_text = fixed_text.replace(old, new)
            fixed_count += 1
    
    # 更新段落
    if fixed_text != original_text:
        para.clear()
        para.add_run(fixed_text)
        print(f"  修正段落 {i}: {original_text[:50]}... -> {fixed_text[:50]}...")

# 保存文档
output_file = "/Volumes/agents/codex/paper2-LSTM量化金融/archive_20260710_1046_chatgpt_review/JFDS_Submission_Package_20260712/paper/main_paper_jfds_READY_v22_numbers_fixed.docx"
doc.save(output_file)

print(f"\n✅ 数字修正完成！")
print(f"修正数量: {fixed_count}")
print(f"保存为: {output_file}")

