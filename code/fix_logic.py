import docx
import re

# Load the document
doc = docx.Document("main_paper_jfds_READY_v14_tables_added.docx")

print("=== 修正逻辑表述 ===\n")

# 需要修正的表述
old_statements = [
    'Source diversity is a necessary condition for successful transfer.',
    'source diversity is a necessary condition for transfer',
    'source diversity is a necessary condition for successful transfer',
    'establish source diversity as a necessary condition',
    'source policy diversity is a necessary condition',
    'source diversity as a necessary condition',
    'source diversity as a transfer prerequisite',
]

new_statements = [
    'Source diversity is associated with successful transfer in our experimental setting.',
    'source diversity is associated with transfer success in our experiments',
    'source diversity is associated with successful transfer in our setting',
    'establish source diversity as a useful screening diagnostic',
    'source policy diversity is a useful screening indicator',
    'source diversity as a practical screening diagnostic',
    'source diversity as a practical screening criterion',
]

fixed_count = 0
for i, para in enumerate(doc.paragraphs):
    original_text = para.text
    fixed_text = original_text
    
    # 修正必要条件表述
    for old, new in zip(old_statements, new_statements):
        if old in fixed_text.lower():
            # 使用正则表达式进行大小写不敏感替换
            pattern = re.compile(re.escape(old), re.IGNORECASE)
            fixed_text = pattern.sub(new, fixed_text)
            fixed_count += 1
    
    # 修正其他过强的表述
    replacements = {
        'demonstrate that source diversity is a necessary condition': 'show that source diversity is associated with transfer success',
        'prove that source diversity is a necessary condition': 'find that source diversity is associated with transfer success',
        'necessary condition for transfer': 'useful indicator for transfer screening',
        'necessary condition for successful transfer': 'useful indicator for transfer success',
    }
    
    for old, new in replacements.items():
        if old in fixed_text.lower():
            pattern = re.compile(re.escape(old), re.IGNORECASE)
            fixed_text = pattern.sub(new, fixed_text)
            fixed_count += 1
    
    # 更新段落
    if fixed_text != original_text:
        para.clear()
        para.add_run(fixed_text)
        print(f"  修复段落 {i}: {original_text[:60]}...")

# 保存文档
output_file = "main_paper_jfds_READY_v15_logic_fixed.docx"
doc.save(output_file)

print(f"\n✅ 逻辑表述修复完成！")
print(f"修复数量: {fixed_count}")
print(f"保存为: {output_file}")

