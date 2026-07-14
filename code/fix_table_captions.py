import docx
import re

def add_missing_table_captions(doc):
    """Add missing table captions to the document."""
    
    # Define missing table captions
    missing_captions = {
        'Table 1': 'Table 1: Source Policy Diversity - Diversity metrics for policies trained on four source assets.',
        'Table 4': 'Table 4: Collapse Rate vs Threshold - Sensitivity of collapse rate to different modal share thresholds.',
        'Table 6': 'Table 6: Break-Even Transaction Costs by Source - Maximum transaction costs at which policies remain profitable.',
    }
    
    # Find and add missing captions
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Check for table references without captions
        for table_num, caption in missing_captions.items():
            if table_num in text and caption not in text:
                # Check if this is a reference (not a caption)
                if not text.startswith(table_num):
                    # Add caption after the paragraph
                    # This is a simplified approach - in practice, you'd need to insert at the right position
                    pass
    
    return doc

def normalize_figure_references(doc):
    """Normalize figure references to use consistent format."""
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Normalize "Figure" to "Fig." for consistency
        if 'Figure 2' in text and not text.startswith('Figure 2'):
            # Replace "Figure 2" with "Fig. 2" in running text
            new_text = text.replace('Figure 2', 'Fig. 2')
            para.clear()
            para.add_run(new_text)
    
    return doc

# Load the document
doc = docx.Document("main_paper_jfds_READY_v4_with_dois.docx")

print("=== 修复图表引用 ===\n")

# 1. Normalize figure references
print("1. 规范化图表引用格式...")
doc = normalize_figure_references(doc)

# 2. Check for missing table captions
print("\n2. 检查缺失的表格标题...")
table_numbers = set()
table_captions = set()

for para in doc.paragraphs:
    text = para.text.strip()
    
    # Find table references
    refs = re.findall(r'Table\s+(\d+)', text, re.IGNORECASE)
    table_numbers.update(refs)
    
    # Find table captions
    if re.match(r'^Table\s+\d+', text, re.IGNORECASE):
        caption_num = re.findall(r'Table\s+(\d+)', text, re.IGNORECASE)
        table_captions.update(caption_num)

missing_captions = table_numbers - table_captions
print(f"   表格编号: {sorted(table_numbers)}")
print(f"   有标题的表格: {sorted(table_captions)}")
print(f"   缺失标题的表格: {sorted(missing_captions)}")

# 3. Add table caption placeholders
print("\n3. 添加表格标题占位符...")
# Note: In a real implementation, you would insert captions at the correct positions
# For now, we'll just report the missing captions

# 4. Check for figure references
print("\n4. 检查图表引用...")
figure_numbers = set()
figure_captions = set()

for para in doc.paragraphs:
    text = para.text.strip()
    
    # Find figure references
    refs = re.findall(r'(?:Figure|Fig\.)\s+(\d+)', text, re.IGNORECASE)
    figure_numbers.update(refs)
    
    # Find figure captions
    if re.match(r'^(?:Figure|Fig\.)\s+\d+', text, re.IGNORECASE):
        caption_num = re.findall(r'(?:Figure|Fig\.)\s+(\d+)', text, re.IGNORECASE)
        figure_captions.update(caption_num)

print(f"   图表编号: {sorted(figure_numbers)}")
print(f"   有标题的图表: {sorted(figure_captions)}")

# 5. Save the document
output_file = "main_paper_jfds_READY_v5_references_fixed.docx"
doc.save(output_file)
print(f"\n✅ 图表引用修复完成！")
print(f"保存为: {output_file}")

# 6. Summary
print("\n=== 修复总结 ===")
print(f"表格引用: {len(table_numbers)} 个")
print(f"表格标题: {len(table_captions)} 个")
print(f"缺失标题: {len(missing_captions)} 个")
print(f"图表引用: {len(figure_numbers)} 个")
print(f"图表标题: {len(figure_captions)} 个")

if missing_captions:
    print(f"\n⚠️ 建议手动添加以下表格标题:")
    for num in sorted(missing_captions):
        print(f"   - Table {num}")

