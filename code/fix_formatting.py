import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_font_for_all_paragraphs(doc, font_name='Times New Roman', font_size=12):
    """Set font for all paragraphs in the document."""
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            # Set font for East Asian text
            r = run._element
            rPr = r.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:eastAsia'), font_name)

def set_line_spacing(doc, spacing=1.5):
    """Set line spacing for all paragraphs."""
    for para in doc.paragraphs:
        para.paragraph_format.line_spacing = spacing

def add_page_numbers(doc):
    """Add page numbers to the document."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        
        # Add page number
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add page number field
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._element.append(fldChar1)
        
        run2 = paragraph.add_run()
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        run2._element.append(instrText)
        
        run3 = paragraph.add_run()
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run3._element.append(fldChar2)

def add_line_numbers(doc):
    """Add line numbers to the document."""
    for section in doc.sections:
        sectPr = section._sectPr
        lnNumType = OxmlElement('w:lnNumType')
        lnNumType.set(qn('w:countBy'), '1')
        lnNumType.set(qn('w:restart'), 'newPage')
        sectPr.append(lnNumType)

# Load the document
doc = docx.Document("main_paper_jfds_READY_v2.docx")

print("正在修复文档格式...")

# 1. 设置字体
print("1. 设置字体为Times New Roman, 12pt...")
set_font_for_all_paragraphs(doc, 'Times New Roman', 12)

# 2. 设置行距
print("2. 设置行距为1.5倍...")
set_line_spacing(doc, 1.5)

# 3. 添加页码
print("3. 添加页码...")
add_page_numbers(doc)

# 4. 添加行号
print("4. 添加行号...")
add_line_numbers(doc)

# Save the document
output_file = "main_paper_jfds_READY_v3_formatted.docx"
doc.save(output_file)
print(f"\n✅ 格式修复完成！")
print(f"保存为: {output_file}")

# Verify the changes
print("\n=== 验证修改 ===")
doc2 = docx.Document(output_file)

# Check fonts
fonts = set()
sizes = set()
for para in doc2.paragraphs:
    for run in para.runs:
        if run.font.name:
            fonts.add(run.font.name)
        if run.font.size:
            sizes.add(run.font.size)

print(f"字体: {fonts}")
print(f"字号: {sizes}")

# Check line spacing
spacings = set()
for para in doc2.paragraphs:
    if para.paragraph_format.line_spacing:
        spacings.add(para.paragraph_format.line_spacing)
print(f"行距: {spacings}")

