import docx
from docx.shared import Pt
import re

# 加载文档
doc = docx.Document("/Volumes/agents/codex/paper2-LSTM量化金融/archive_20260710_1046_chatgpt_review/JFDS_Submission_Package_20260712/paper/论文_公式恢复与图表美化_RHETORIC.docx")

print("=== 全面修复数字一致性问题 ===\n")

# 首先找出所有包含234的段落
print("1. 查找所有包含'234'的段落:")
for i, para in enumerate(doc.paragraphs):
    if '234' in para.text:
        print(f"   段落 {i}: {para.text[:100]}...")

print("\n2. 查找所有包含'39 target'的段落:")
for i, para in enumerate(doc.paragraphs):
    if '39 target' in para.text.lower() or '39targets' in para.text.lower():
        print(f"   段落 {i}: {para.text[:100]}...")

print("\n3. 查找所有包含'6×39'或'6 × 39'的段落:")
for i, para in enumerate(doc.paragraphs):
    if '6×39' in para.text or '6 × 39' in para.text:
        print(f"   段落 {i}: {para.text[:100]}...")

print("\n4. 查找Table 2相关内容:")
for i, para in enumerate(doc.paragraphs):
    if 'Table 2' in para.text or '228' in para.text:
        print(f"   段落 {i}: {para.text[:100]}...")

