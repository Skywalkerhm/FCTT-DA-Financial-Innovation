import docx
from docx.shared import Pt
import re

def read_markdown_file(filepath):
    """读取markdown文件"""
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()

def replace_section_in_doc(doc, section_keyword, new_content):
    """替换文档中的章节内容"""
    start_idx = None
    end_idx = None
    
    # 找到章节开始位置
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if section_keyword in text:
            start_idx = i
            break
    
    if start_idx is None:
        print(f"  ❌ 未找到 {section_keyword}")
        return False
    
    # 找到章节结束位置（下一个同级标题）
    for i in range(start_idx + 1, len(doc.paragraphs)):
        text = doc.paragraphs[i].text.strip()
        style = doc.paragraphs[i].style.name if doc.paragraphs[i].style else ""
        
        # 检查是否是下一个主要章节
        if 'Heading 1' in style and text and text != section_keyword:
            end_idx = i
            break
    
    if end_idx is None:
        end_idx = len(doc.paragraphs)
    
    print(f"  替换段落 {start_idx} 到 {end_idx-1}")
    
    # 删除原有内容（从后往前删除）
    for i in range(end_idx - 1, start_idx, -1):
        p = doc.paragraphs[i]
        p._element.getparent().remove(p._element)
    
    # 添加新内容
    lines = new_content.strip().split('\n')
    current_para = doc.paragraphs[start_idx]
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        new_para = doc.add_paragraph()
        
        # 根据markdown格式设置样式
        if line.startswith('## '):
            new_para.style = doc.styles['Heading 1']
            new_para.add_run(line[3:])
        elif line.startswith('### '):
            new_para.style = doc.styles['Heading 2']
            new_para.add_run(line[4:])
        elif line.startswith('**') and line.endswith('**'):
            run = new_para.add_run(line[2:-2])
            run.bold = True
        elif line.startswith('*') and line.endswith('*'):
            run = new_para.add_run(line[1:-1])
            run.italic = True
        elif line.startswith('- '):
            new_para.style = doc.styles['List Bullet']
            new_para.add_run(line[2:])
        elif re.match(r'^\d+\.\s+', line):
            new_para.style = doc.styles['List Number']
            new_para.add_run(line)
        else:
            new_para.add_run(line)
        
        # 在当前位置后插入
        current_para._element.addnext(new_para._element)
        current_para = new_para
    
    return True

# 加载文档
doc = docx.Document("/Volumes/agents/codex/paper2-LSTM量化金融/archive_20260710_1046_chatgpt_review/JFDS_Submission_Package_20260712/paper/main_paper_jfds_FINAL_COMPLETE.docx")

print("=== 修正理论部分 ===\n")

# 读取修正后的理论内容
new_theory = read_markdown_file("/Volumes/agents/codex/paper2-LSTM量化金融/archive_20260710_1046_chatgpt_review/JFDS_Submission_Package_20260712/paper/fix_theoretical_propositions.md")

# 替换5. Theoretical Analysis章节
print("替换5. Theoretical Analysis章节...")
success = replace_section_in_doc(doc, '5. Theoretical Analysis', new_theory)

if success:
    # 保存文档
    output_file = "/Volumes/agents/codex/paper2-LSTM量化金融/archive_20260710_1046_chatgpt_review/JFDS_Submission_Package_20260712/paper/main_paper_jfds_READY_v24_theory_fixed.docx"
    doc.save(output_file)
    print(f"\n✅ 理论部分修正完成！")
    print(f"保存为: {output_file}")
else:
    print("\n❌ 修正失败")

