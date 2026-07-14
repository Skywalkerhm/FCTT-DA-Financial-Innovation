import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def fix_math_symbols(text):
    """Fix mathematical symbols in text."""
    # Fix transpose symbols
    text = text.replace('^T', '^\\top')
    
    # Fix common patterns
    replacements = {
        'θ_k^T': '$\\theta_k^\\top$',
        'θ_j^T': '$\\theta_j^\\top$',
        'θ_a^T': '$\\theta_a^\\top$',
        'θ_k^\\top': '$\\theta_k^\\top$',
        'θ_j^\\top': '$\\theta_j^\\top$',
        'θ_a^\\top': '$\\theta_a^\\top$',
    }
    
    for old, new in replacements.items():
        text = text.replace(old, new)
    
    return text

def fix_paragraph_formatting(para):
    """Fix paragraph formatting."""
    # Set font for all runs
    for run in para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

def add_equation_numbering(doc):
    """Add equation numbering to display equations."""
    equation_count = 0
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # Check if this is a display equation
        if text.startswith('$$') and text.endswith('$$'):
            equation_count += 1
            
            # Add equation number
            new_text = f"{text}  ({equation_count})"
            para.clear()
            para.add_run(new_text)
            
            # Center align
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return equation_count

def fix_table_formatting(doc):
    """Fix table formatting."""
    for table in doc.tables:
        # Set font for all cells
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(10)
                    
                    # Center align cells
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def fix_heading_formatting(doc):
    """Fix heading formatting."""
    for para in doc.paragraphs:
        style = para.style.name if para.style else ""
        
        if 'Heading 1' in style:
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(16)
                run.bold = True
        elif 'Heading 2' in style:
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                run.bold = True
        elif 'Heading 3' in style:
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.bold = True

# Load the document
doc = docx.Document("main_paper_jfds_READY_v19_ablation.docx")

print("=== 完善公式和格式 ===\n")

# 1. 修复数学符号
print("1. 修复数学符号...")
fixed_count = 0
for i, para in enumerate(doc.paragraphs):
    original_text = para.text
    fixed_text = fix_math_symbols(original_text)
    
    if fixed_text != original_text:
        para.clear()
        para.add_run(fixed_text)
        fixed_count += 1

print(f"   修复了 {fixed_count} 个段落的数学符号")

# 2. 添加公式编号
print("\n2. 添加公式编号...")
equation_count = add_equation_numbering(doc)
print(f"   添加了 {equation_count} 个公式编号")

# 3. 修复段落格式
print("\n3. 修复段落格式...")
for i, para in enumerate(doc.paragraphs):
    fix_paragraph_formatting(para)

print("   修复了所有段落的字体格式")

# 4. 修复表格格式
print("\n4. 修复表格格式...")
fix_table_formatting(doc)
print("   修复了所有表格的格式")

# 5. 修复标题格式
print("\n5. 修复标题格式...")
fix_heading_formatting(doc)
print("   修复了所有标题的格式")

# 6. 设置页面格式
print("\n6. 设置页面格式...")
for section in doc.sections:
    # 设置页边距
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # 设置页面大小
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)

print("   设置了页面格式")

# 保存文档
output_file = "main_paper_jfds_READY_v20_formatted.docx"
doc.save(output_file)

print(f"\n✅ 公式和格式完善完成！")
print(f"保存为: {output_file}")

# 统计信息
total_words = sum(len(re.findall(r'\w+', para.text)) for para in doc.paragraphs)
print(f"文档总字数: {total_words}")
print(f"公式数量: {equation_count}")
print(f"表格数量: {len(doc.tables)}")

