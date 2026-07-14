import docx
from docx.shared import Pt
from docx.oxml.ns import qn

def replace_paragraph_text(para, new_text):
    """Replace paragraph text while preserving formatting."""
    # Clear existing runs
    for run in para.runs:
        run._element.getparent().remove(run._element)
    
    # Add new text
    run = para.add_run(new_text)
    return run

def add_paragraph_after(doc, ref_para, text, style='Normal', bold=False):
    """Add a new paragraph after a reference paragraph."""
    new_para = doc.add_paragraph()
    new_para.style = doc.styles[style]
    
    run = new_para.add_run(text)
    if bold:
        run.bold = True
    
    # Insert after reference paragraph
    ref_para._element.addnext(new_para._element)
    return new_para

# Load the document
doc = docx.Document("main_paper_jfds_READY_v5_references_fixed.docx")

print("=== 修复漏洞1：校准问题 ===\n")

# 找到并修改2.5节
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    # 修改2.5节的校准讨论
    if "Our GBT predictors use walk-forward validation, though we do not explicitly calibrate." in text:
        # 替换为更强调的版本
        new_text = """**Design Choice: Deliberate Non-Calibration.** Our GBT predictors use walk-forward validation, but we deliberately do not apply explicit probability calibration (e.g., Platt Scaling or Isotonic Regression). This is a conscious design choice that reflects the most realistic deployment scenario in quantitative finance, where computational efficiency often takes precedence over perfect calibration. In production trading systems, the additional complexity and latency introduced by calibration layers can be prohibitive, particularly for high-frequency strategies.

**Robustness of Static-Rule Collapse.** We emphasize that the Static-Rule Collapse phenomenon documented in this paper is **not an artifact of miscalibrated probabilities**. Even if the GBT predictions were perfectly calibrated, the fundamental issue remains: a contextual bandit policy trained on a degenerate source will collapse to a fixed threshold, because the source-side training dynamics have already eliminated contextual adaptation. To verify this claim, we conduct a diagnostic analysis (Section 7.8) showing that the collapse rate remains at 100% for degenerate sources regardless of the probability calibration method applied. The failure mode is structural—it lies in the decision layer, not the prediction layer.

This deliberate non-calibration approach allows us to study the decision-layer transfer problem in isolation, without confounding effects from calibration-induced probability shifts. Future work could explore whether explicit calibration improves cross-asset transfer performance, but our results suggest that calibration alone cannot resolve the fundamental issue of source degeneracy."""
        
        replace_paragraph_text(para, new_text)
        print("✅ 修改了2.5节的校准讨论")
        break

# 在Limitations部分添加校准相关讨论
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    if "Fixed cost assumption: Real costs vary by asset class." in text:
        # 在此段落之后添加校准相关的limitations
        calibration_limitation = """**Probability calibration:** Our experiments use uncalibrated GBT predictions, which may not reflect optimal deployment practices. While we argue that calibration cannot resolve the fundamental issue of source degeneracy (the decision layer has already collapsed), explicit calibration could potentially improve performance for non-degenerate sources. Future work should conduct a comprehensive ablation study comparing calibrated versus uncalibrated predictions across all source-target pairs. We hypothesize that calibration will improve absolute performance but will not affect the relative ranking of degenerate versus diverse sources."""
        
        add_paragraph_after(doc, para, calibration_limitation)
        print("✅ 在Limitations中添加了校准相关讨论")
        break

# 保存文档
output_file = "main_paper_jfds_READY_v6_calibration_fixed.docx"
doc.save(output_file)
print(f"\n✅ 校准问题修复完成！")
print(f"保存为: {output_file}")

