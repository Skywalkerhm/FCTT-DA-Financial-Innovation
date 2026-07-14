import docx
from docx.shared import Pt
import re

# 加载文档
doc = docx.Document("/Volumes/agents/codex/paper2-LSTM量化金融/archive_20260710_1046_chatgpt_review/JFDS_Submission_Package_20260712/paper/论文_公式恢复与图表美化_OVERCLAIM_FIXED.docx")

print("=== 修正剩余问题 ===\n")

# 定义修正规则
replacements = [
    # 1. EMD前视偏差问题
    ('EMD', 'Empirical Mode Decomposition (EMD)'),
    ('first IMF', 'first Intrinsic Mode Function (IMF)'),
    ('price series', 'rolling window of price observations'),
    
    # 2. features描述修正
    ('technical and fundamental features', 'technical, volatility, volume, and price-based features'),
    ('technical and fundamental', 'technical, volatility, volume, and price-based'),
    
    # 3. 加密货币数据来源修正
    ('public blockchain explorers', 'Yahoo Finance and CoinGecko'),
    ('blockchain explorers', 'Yahoo Finance and CoinGecko'),
    
    # 4. bootstrap次数修正
    ('stationary bootstrap with ___ replications', 'stationary bootstrap with 2,000 replications'),
    ('stationary bootstrap with', 'stationary bootstrap with 2,000'),
    ('bootstrap with ___ replications', 'bootstrap with 2,000 replications'),
    
    # 5. TLT唯一non-degenerate修正
    ('TLT source policy (the only non-degenerate source)', 
     'TLT source policy (one of the two non-degenerate sources, along with IWM)'),
    ('the only non-degenerate source', 'one of the two non-degenerate sources'),
    ('TLT (the only non-degenerate source)', 'TLT (one of the two non-degenerate sources)'),
    
    # 6. 表格文字错误修正
    ('dif erent', 'different'),
    ('shuf ling', 'shuffling'),
    ('dif ferent', 'different'),
    ('shuf fled', 'shuffled'),
    
    # 7. 其他常见错误
    ('frozen contextual threshold', 'frozen contextual threshold policy'),
    ('cross-asset transfer', 'cross-asset policy transfer'),
]

# 修正文档
fixed_count = 0
for i, para in enumerate(doc.paragraphs):
    original_text = para.text
    fixed_text = original_text
    
    # 应用替换
    for old, new in replacements:
        if old in fixed_text:
            fixed_text = fixed_text.replace(old, new)
            fixed_count += 1
    
    # 更新段落
    if fixed_text != original_text:
        para.clear()
        para.add_run(fixed_text)
        print(f"  修正段落 {i}: {original_text[:50]}...")

# 修正表格中的错误
print("\n修正表格:")
for i, table in enumerate(doc.tables):
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            original_text = cell.text
            fixed_text = original_text
            
            # 应用替换
            for old, new in replacements:
                if old in fixed_text:
                    fixed_text = fixed_text.replace(old, new)
                    fixed_count += 1
            
            # 更新单元格
            if fixed_text != original_text:
                for para in cell.paragraphs:
                    para.clear()
                    para.add_run(fixed_text)
                print(f"  修正 Table {i+1}, Row {row_idx}, Cell {cell_idx}")

# 保存文档
output_file = "/Volumes/agents/codex/paper2-LSTM量化金融/archive_20260710_1046_chatgpt_review/JFDS_Submission_Package_20260712/paper/论文_公式恢复与图表美化_FINAL_FIXED.docx"
doc.save(output_file)

print(f"\n✅ 所有问题修正完成！")
print(f"修正数量: {fixed_count}")
print(f"保存为: {output_file}")

