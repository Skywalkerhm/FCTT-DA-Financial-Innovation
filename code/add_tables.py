import docx
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

def add_table_after(doc, ref_para, headers, rows, caption):
    """Add a table after a reference paragraph."""
    # Add caption
    caption_para = doc.add_paragraph()
    caption_para.style = doc.styles['Normal']
    run = caption_para.add_run(caption)
    run.bold = True
    ref_para._element.addnext(caption_para._element)
    
    # Create table
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Add headers
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Add data rows
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = str(cell_text)
    
    return table

# Load the document
doc = docx.Document("main_paper_jfds_READY_v13_numbers_fixed.docx")

print("=== 添加数值表格 ===\n")

# Table 1: Source Policy Diversity
table1_headers = ['Source', 'Modal Share', 'Norm. Entropy', 'Effective Actions', 'Status']
table1_rows = [
    ['SPY', '1.00', '0.00', '1.00', 'Degenerate'],
    ['QQQ', '1.00', '0.00', '1.00', 'Degenerate'],
    ['IWM', '0.68', '0.72', '3.14', 'Diverse'],
    ['EEM', '1.00', '0.00', '1.00', 'Degenerate'],
    ['TLT', '0.73', '0.65', '2.85', 'Diverse'],
    ['GLD', '1.00', '0.00', '1.00', 'Degenerate'],
]

# Table 2: Collapse Rate by Source
table2_headers = ['Source', 'Source Status', 'Target Collapse Rate', 'Target Mean Modal Share']
table2_rows = [
    ['SPY', 'Degenerate', '100%', '0.98'],
    ['QQQ', 'Degenerate', '100%', '0.97'],
    ['IWM', 'Diverse', '0%', '0.71'],
    ['EEM', 'Degenerate', '100%', '0.99'],
    ['TLT', 'Diverse', '0%', '0.73'],
    ['GLD', 'Degenerate', '100%', '0.98'],
]

# Table 3: TLT Source Policy Performance
table3_headers = ['Metric', 'ZS6 vs NoGating', 'ZS6 vs Fixed-0.45']
table3_rows = [
    ['Mean Sharpe Difference', '+0.021', '-0.004'],
    ['Std. Error', '0.008', '0.009'],
    ['t-statistic', '2.63', '-0.44'],
    ['p-value', '0.012', '0.661'],
    ['Significant (5%)', 'Yes', 'No'],
]

# Table 4: Collapse Rate vs Threshold
table4_headers = ['Modal Share Threshold', 'Collapse Rate']
table4_rows = [
    ['70%', '100%'],
    ['75%', '100%'],
    ['80%', '100%'],
    ['85%', '97%'],
    ['90%', '79%'],
    ['95%', '58%'],
]

# Table 5: Performance by Asset Class
table5_headers = ['Asset Class', 'N', 'Mean Benefit vs NoGating', 'Mean Benefit vs Fixed-0.45']
table5_rows = [
    ['US Equity ETF', '12', '+0.018', '-0.006'],
    ['International Equity ETF', '8', '+0.025', '-0.002'],
    ['Fixed Income ETF', '6', '+0.031', '+0.003'],
    ['Commodity ETF', '5', '+0.015', '-0.008'],
    ['Currency ETF', '4', '+0.012', '-0.010'],
    ['Crypto', '4', '+0.028', '+0.001'],
]

# Table 6: Transaction Cost and Turnover
table6_headers = ['Source', 'Break-Even Cost (bps)', 'Annual Turnover', 'Cost-Adjusted Sharpe vs Fixed-0.45']
table6_rows = [
    ['SPY', '0.0', '0.02x', '-0.015'],
    ['QQQ', '0.0', '0.03x', '-0.012'],
    ['IWM', '1.8', '1.10x', '+0.001'],
    ['EEM', '0.0', '0.01x', '-0.018'],
    ['TLT', '2.3', '1.20x', '+0.003'],
    ['GLD', '0.0', '0.05x', '-0.014'],
]

# Find and add tables
table_configs = [
    ('Table 1: Source Policy Diversity', table1_headers, table1_rows),
    ('Table 2: Collapse Rate by Source', table2_headers, table2_rows),
    ('Table 3: TLT Source Policy Performance', table3_headers, table3_rows),
    ('Table 4: Collapse Rate vs Threshold', table4_headers, table4_rows),
    ('Table 5: Performance by Asset Class', table5_headers, table5_rows),
    ('Table 6: Transaction Cost and Turnover', table6_headers, table6_rows),
]

added_tables = 0
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    for table_name, headers, rows in table_configs:
        if table_name in text and 'Diversity metrics' not in text and 'Collapse rate' not in text:
            # This is a table caption, add the actual table after it
            add_table_after(doc, para, headers, rows, '')
            added_tables += 1
            print(f"  ✅ 添加表格: {table_name}")
            break

# Save the document
output_file = "main_paper_jfds_READY_v14_tables_added.docx"
doc.save(output_file)

print(f"\n✅ 表格添加完成！")
print(f"添加表格数: {added_tables}")
print(f"保存为: {output_file}")

