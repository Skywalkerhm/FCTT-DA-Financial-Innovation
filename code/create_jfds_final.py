from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

# Create a new document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Set margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Title
title = doc.add_heading('Parameters Transfer, Policies May Not: Diagnosing Static-Rule Collapse in Cross-Asset Decision Gating', 0)

# Author and date
doc.add_paragraph('Min Huang')
doc.add_paragraph('July 2026')
doc.add_paragraph('Target: The Journal of Financial Data Science')

# Abstract
doc.add_heading('Abstract', level=1)
abstract_text = """We identify Static-Rule Collapse, a silent failure mode in cross-asset decision-layer transfer where a frozen contextual policy degenerates into a fixed probability threshold on out-of-distribution targets. This failure mode is undetectable by standard backtests, as the collapsed policy may still outperform naive baselines while losing all contextual adaptation capability.

Across 234 source-target pairs (6 sources × 39 targets), we establish an empirical taxonomy: degenerate sources (SPY, QQQ, EEM, GLD) trigger 100% target collapse; diverse sources (TLT, IWM) maintain contextual adaptation. We prove that source diversity is a necessary condition for transfer and provide sufficient conditions linking source-side training dynamics to target-side failure.

We contribute FCTT-DA (Frozen Contextual Threshold Transfer with Degeneracy Audit), a deployable framework with source diversity pre-checks, behavioral-equivalent baseline comparison, transaction cost sensitivity analysis, and an action-region geometry diagnostic. The framework detects collapse before deployment, preventing futile transfers and saving computational resources.

Our findings have immediate implications for quantitative trading firms deploying machine learning policies across multiple assets. We demonstrate that the common practice of transferring policies from single assets is fundamentally flawed, and propose source diversity as a prerequisite for successful cross-asset transfer."""
doc.add_paragraph(abstract_text)

# Keywords (adjusted to 5-6)
doc.add_paragraph('Keywords: contextual bandits, policy transfer, decision gating, model audit, static-rule collapse, cross-asset learning')

# 1. Introduction
doc.add_heading('1. Introduction', level=1)

doc.add_heading('1.1 The operational appeal of decision-layer reuse', level=2)
intro_text = """In quantitative asset management, the standard workflow separates prediction from decision. A machine-learning model for return forecasting is trained once and reused across multiple assets, exploiting the statistical similarity of financial time series. The decision layer—typically a threshold on predicted probabilities—is often transferred alongside the prediction model, on the assumption that a threshold that works on one asset should work on others.

Consider a concrete example: a quantitative team develops a decision policy for SPDR S&P 500 ETF (SPY) that achieves a Sharpe ratio of 0.45 over five years of walk-forward validation. The natural next step is to deploy this policy on iShares Russell 2000 ETF (IWM) or Invesco QQQ Trust (QQQ). Standard practice compares the transferred policy against a static 0.50 cutoff. If the policy achieves a Sharpe of 0.42 on IWM versus 0.38 for the static cutoff, the team concludes that contextual transfer 'works.' But this conclusion may be entirely spurious."""
doc.add_paragraph(intro_text)

doc.add_heading('1.2 The diagnostic gap', level=2)
diagnostic_text = """Standard backtests compare a learned policy against a conventional cutoff (typically 0.50). But these comparisons cannot distinguish between two fundamentally different scenarios:

1. Contextual transfer: The policy uses market-state information to select different thresholds at different times.
2. Static-rule collapse: The policy consistently selects the same threshold, behaving identically to a simple static rule.

When scenario (2) obtains, apparent improvements over 0.50 are merely the consequence of a different (but still fixed) probability cutoff—say, 0.45 instead of 0.50. The contextual policy has not learned to adapt; it has merely discovered a better static threshold through source-side optimization.

The diagnostic gap we identify is subtle but consequential. When a contextual bandit policy collapses to a fixed threshold, it loses its ability to adapt to changing market conditions. In volatile markets, a fixed threshold may be too aggressive; in calm markets, too conservative. The policy's apparent success on the source asset masks its fundamental failure to learn contextual adaptation."""
doc.add_paragraph(diagnostic_text)

doc.add_heading('1.3 Contributions', level=2)
contributions_text = """Source diversity as a transfer prerequisite. We establish that source policy diversity during training is a necessary condition for successful cross-asset transfer. When the source policy degenerates to a static rule, the target policy inherits this degeneracy, regardless of the target's market dynamics.

FCTT-DA audit framework. We propose a modular architecture with pre-deployment diagnostics that detect source-side degeneracy before transfer. The framework includes normalized action entropy, modal-action share, and effective actions as diversity metrics, with a rejection rule based on source modal share.

Multi-source empirical evidence. We evaluate 234 source-target pairs across 6 asset classes, with paired bootstrap inference and false discovery rate control. The results demonstrate that degenerate sources trigger 100% target collapse, while diverse sources maintain contextual adaptation.

Theoretical conditions. We provide sufficient conditions linking source diversity to target collapse based on decision-region geometry and behavioral equivalence. These conditions yield testable predictions about when transfer will fail.

Our contributions can be summarized as follows. First, we formally define Static-Rule Collapse and demonstrate that it is a pervasive failure mode in cross-asset policy transfer. Second, we propose FCTT-DA as a practical diagnostic framework for detecting collapse before deployment. Third, we establish source diversity as a necessary condition for successful transfer, with implications for how practitioners select source assets. Fourth, we provide theoretical foundations linking source-side training dynamics to target-side failure."""
doc.add_paragraph(contributions_text)

# Save the document
doc.save("main_paper_jfds_final_v2.docx")
print("Created initial document structure")
