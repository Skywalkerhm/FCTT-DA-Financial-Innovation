import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 加载文档
doc = docx.Document("/Volumes/agents/codex/paper2-LSTM量化金融/archive_20260710_1046_chatgpt_review/JFDS_Submission_Package_20260712/paper/论文_公式恢复与图表美化_FINAL.docx")

print("=== 添加页码 ===\n")

for section in doc.sections:
    # 获取或创建footer
    footer = section.footer
    footer.is_linked_to_previous = False
    
    # 清除现有内容
    for para in footer.paragraphs:
        para.clear()
    
    # 创建页码段落
    if footer.paragraphs:
        paragraph = footer.paragraphs[0]
    else:
        paragraph = footer.add_paragraph()
    
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加页码字段
    run = paragraph.add_run()
    
    # 开始字段
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)
    
    # 字段指令
    run2 = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run2._element.append(instrText)
    
    # 结束字段
    run3 = paragraph.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run3._element.append(fldChar2)
    
    print(f"  ✅ 为section添加页码")

# 保存文档
output_file = "/Volumes/agents/codex/paper2-LSTM量化金融/archive_20260710_1046_chatgpt_review/JFDS_Submission_Package_20260712/paper/论文_公式恢复与图表美化_FINAL.docx"
doc.save(output_file)

print(f"\n✅ 页码添加完成！")

# 验证
doc2 = docx.Document(output_file)
has_page_numbers = False
for section in doc2.sections:
    footer = section.footer
    if footer.paragraphs:
        for para in footer.paragraphs:
            if 'PAGE' in para.text or para.text.strip():
                has_page_numbers = True
                break

print(f"页码验证: {'有' if has_page_numbers else '无'}")

