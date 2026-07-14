import docx
import re

def fix_math_symbols(text):
    """Fix mathematical symbols in text."""
    # Fix transpose symbols: ^T -> ^\top
    text = re.sub(r'\^T\b', r'^\\top', text)
    
    # Fix inline math that's not properly formatted
    # Pattern: θ_k^T c -> $\theta_k^\top \mathbf{c}$
    text = re.sub(r'θ_k\^T c', r'$\\theta_k^\\top \\mathbf{c}$', text)
    text = re.sub(r'θ_k\^\\top c', r'$\\theta_k^\\top \\mathbf{c}$', text)
    
    # Fix other common patterns
    text = text.replace('θ_k^T', '$\\theta_k^\\top$')
    text = text.replace('θ_j^T', '$\\theta_j^\\top$')
    text = text.replace('θ_a^T', '$\\theta_a^\\top$')
    
    # Fix π_S
    text = text.replace('π_S', '$\\pi_S$')
    
    # Fix ∈
    text = text.replace('∈', '$\\in$')
    
    # Fix ℝ
    text = text.replace('ℝ', '$\\mathbb{R}$')
    
    # Fix ×
    text = text.replace('×', '$\\times$')
    
    return text

def fix_paragraph_math(para):
    """Fix math symbols in a paragraph."""
    original_text = para.text
    fixed_text = fix_math_symbols(original_text)
    
    if fixed_text != original_text:
        # Clear and rewrite
        para.clear()
        para.add_run(fixed_text)
        return True
    return False

# Load the document
doc = docx.Document("main_paper_jfds_READY_v9_reproducibility.docx")

print("=== 修复公式符号 ===\n")

# Find and fix Appendix F
in_appendix_f = False
fixed_count = 0

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    if "Appendix F: Context Cloud Geometry" in text:
        in_appendix_f = True
    
    if in_appendix_f:
        if text.startswith("Appendix") and "F" not in text:
            break
        
        # Fix math symbols
        if fix_paragraph_math(para):
            fixed_count += 1
            print(f"  修复段落 {i}: {text[:80]}...")

# Also fix main paper sections that reference these formulas
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    # Fix sections that reference Appendix F formulas
    if 'theta_k' in text.lower() or 'θ_k' in text:
        if fix_paragraph_math(para):
            fixed_count += 1
            print(f"  修复段落 {i}: {text[:80]}...")

# Save the document
output_file = "main_paper_jfds_READY_v10_math_fixed.docx"
doc.save(output_file)

print(f"\n✅ 公式符号修复完成！")
print(f"修复段落数: {fixed_count}")
print(f"保存为: {output_file}")

