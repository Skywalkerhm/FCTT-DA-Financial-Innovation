import docx
from docx.shared import Pt
import re

def fix_text(text, old_patterns, new_patterns):
    """替换文本中的模式"""
    for old, new in zip(old_patterns, new_patterns):
        text = text.replace(old, new)
    return text

# 加载文档
doc = docx.Document("/Volumes/agents/codex/paper2-LSTM量化金融/archive_20260710_1046_chatgpt_review/JFDS_Submission_Package_20260712/paper/main_paper_jfds_FINAL_THEORY_FIXED.docx")

print("=== 修正逻辑矛盾和核心叙事 ===\n")

# 需要修正的模式
old_patterns = [
    # 逻辑矛盾
    'Source degeneracy is a necessary condition for severe target collapse',
    'source degeneracy is not sufficient for target collapse',
    'Source diversity is a necessary condition for successful transfer',
    'source diversity is a necessary condition for transfer',
    'establish source diversity as a necessary condition',
    
    # 核心叙事问题
    'a contextual policy degenerates on out-of-distribution targets',
    'degenerates on out-of-distribution targets',
    'cross-asset transfer causes collapse',
    'transfer causes collapse',
    
    # 术语问题
    'Static-Rule Collapse',
    'static-rule collapse',
    
    # 审计重点
    'FCTT-DA detects collapse before deployment',
    'detects collapse before deployment',
]

new_patterns = [
    # 逻辑修正
    'In our empirical sample, source degeneracy perfectly predicts target collapse',
    'source degeneracy is not logically necessary for target collapse in general',
    'In our sample, source diversity is associated with target diversity',
    'source diversity is associated with transfer success in our experiments',
    'identify source diversity as a useful screening diagnostic',
    
    # 核心叙事修正
    'a policy that is already degenerate on the source asset propagates this degeneracy to target assets',
    'propagates source-side degeneracy to target assets',
    'cross-asset transfer propagates source-side degeneracy',
    'transfer propagates degeneracy',
    
    # 术语修正
    'Source Static Degeneracy',
    'source static degeneracy',
    
    # 审计重点修正
    'FCTT-DA detects source-side degeneracy before transfer',
    'detects source-side degeneracy before transfer',
]

# 修正文档
fixed_count = 0
for i, para in enumerate(doc.paragraphs):
    original_text = para.text
    fixed_text = original_text
    
    # 替换模式
    for old, new in zip(old_patterns, new_patterns):
        if old in fixed_text:
            fixed_text = fixed_text.replace(old, new)
            fixed_count += 1
    
    # 更新段落
    if fixed_text != original_text:
        para.clear()
        para.add_run(fixed_text)
        print(f"  修正段落 {i}: {original_text[:50]}...")

# 保存文档
output_file = "/Volumes/agents/codex/paper2-LSTM量化金融/archive_20260710_1046_chatgpt_review/JFDS_Submission_Package_20260712/paper/main_paper_jfds_READY_v25_logic_fixed.docx"
doc.save(output_file)

print(f"\n✅ 逻辑矛盾修正完成！")
print(f"修正数量: {fixed_count}")
print(f"保存为: {output_file}")

