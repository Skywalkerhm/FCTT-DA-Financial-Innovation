import docx
import re

# Load the document
doc = docx.Document("main_paper_jfds_READY_v17_methods_updated.docx")

print("=== 调整论文主张 ===\n")

# 需要修正的过强表述
claim_adjustments = {
    # 过强的理论主张
    'demonstrate that it is pervasive': 'show that it occurs in our experimental setting',
    'demonstrate that source diversity is a necessary condition': 'show that source diversity is associated with transfer success',
    'prove that source diversity is a necessary condition': 'find that source diversity is associated with transfer success',
    'establish source diversity as a necessary condition': 'identify source diversity as a useful screening indicator',
    'This represents a paradigm shift': 'This suggests a useful perspective',
    'paradigm shift in model audit': 'useful perspective for model audit',
    'first practical tool': 'practical tool',
    'first paper to': 'paper that',
    'To our knowledge, this is the first': 'This paper',
    
    # 营销式表达
    'immediate practical impact': 'practical implications',
    'actionable value': 'practical value',
    'critical gap': 'gap',
    'urgent need': 'need',
    'fundamental question': 'question',
    'paradigm shift': 'perspective shift',
    'novel failure mode': 'failure mode',
    'silent failure mode': 'failure mode',
    
    # 过度承诺
    'provides the most comprehensive study': 'provides a comprehensive study',
    'most comprehensive study': 'comprehensive study',
    'provides immediate value': 'provides value',
    'direct cost savings': 'cost savings',
    'saves computational resources': 'reduces computational burden',
}

# 修正结论部分
print("1. 修正结论部分...")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '9. Conclusion' in text or 'Conclusion' in text:
        # 找到结论部分，修正其内容
        for j in range(i, min(i+20, len(doc.paragraphs))):
            original_text = doc.paragraphs[j].text
            fixed_text = original_text
            
            for old, new in claim_adjustments.items():
                if old in fixed_text:
                    pattern = re.compile(re.escape(old), re.IGNORECASE)
                    fixed_text = pattern.sub(new, fixed_text)
            
            if fixed_text != original_text:
                doc.paragraphs[j].clear()
                doc.paragraphs[j].add_run(fixed_text)
                print(f"   修正段落 {j}")

# 修正摘要部分
print("\n2. 修正摘要部分...")
for i, para in enumerate(doc.paragraphs[:30]):
    original_text = para.text
    fixed_text = original_text
    
    for old, new in claim_adjustments.items():
        if old in fixed_text:
            pattern = re.compile(re.escape(old), re.IGNORECASE)
            fixed_text = pattern.sub(new, fixed_text)
    
    if fixed_text != original_text:
        para.clear()
        para.add_run(fixed_text)
        print(f"   修正段落 {i}")

# 修正Introduction部分
print("\n3. 修正Introduction部分...")
in_intro = False
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '1. Introduction' in text:
        in_intro = True
    if in_intro and '2. Related Work' in text:
        break
    
    if in_intro:
        original_text = para.text
        fixed_text = original_text
        
        for old, new in claim_adjustments.items():
            if old in fixed_text:
                pattern = re.compile(re.escape(old), re.IGNORECASE)
                fixed_text = pattern.sub(new, fixed_text)
        
        if fixed_text != original_text:
            para.clear()
            para.add_run(fixed_text)
            print(f"   修正段落 {i}")

# 保存文档
output_file = "main_paper_jfds_READY_v18_claims_adjusted.docx"
doc.save(output_file)

print(f"\n✅ 论文主张调整完成！")
print(f"保存为: {output_file}")

