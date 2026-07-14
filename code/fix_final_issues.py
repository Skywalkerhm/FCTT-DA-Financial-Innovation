import docx
import re

def add_table_caption(doc, table_num, caption_text, after_para):
    """Add a table caption after a paragraph."""
    new_para = doc.add_paragraph()
    new_para.style = doc.styles['Normal']
    run = new_para.add_run(caption_text)
    run.bold = True
    after_para._element.addnext(new_para._element)
    return new_para

# Load the document
doc = docx.Document("main_paper_jfds_READY_v10_math_fixed.docx")

print("=== 修复最终问题 ===\n")

# 1. 添加缺失的表格标题
print("1. 添加缺失的表格标题...")

# 定义缺失的表格标题
missing_captions = {
    '1': 'Table 1: Source Policy Diversity - Diversity metrics for policies trained on four source assets (SPY, QQQ, EEM, TLT).',
    '4': 'Table 4: Collapse Rate vs Threshold - Sensitivity of collapse rate to different modal share thresholds (70%-95%).',
    '6': 'Table 6: Transaction Cost and Turnover Summary - Break-even costs and annual turnover by source asset.'
}

# 查找表格引用位置
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    # 查找Table 1引用
    if 'Table 1' in text and 'Source Policy Diversity' not in text:
        # 在此段落之后添加标题
        new_para = doc.add_paragraph()
        new_para.style = doc.styles['Normal']
        run = new_para.add_run(missing_captions['1'])
        run.bold = True
        para._element.addnext(new_para._element)
        print(f"   ✅ 添加Table 1标题 (段落 {i})")
        break

# 2. 修复重复的附录D
print("\n2. 修复重复的附录D...")
appendix_d_count = 0
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if 'Appendix D: Reproducibility' in text:
        appendix_d_count += 1
        if appendix_d_count > 1:
            # 删除重复的附录D标题
            para.clear()
            print(f"   ✅ 删除重复的附录D标题 (段落 {i})")

# 3. 检查图表标题
print("\n3. 检查图表标题...")
figure_captions_found = 0
for para in doc.paragraphs:
    text = para.text.strip()
    if re.match(r'^Figure\s+\d+', text, re.IGNORECASE):
        figure_captions_found += 1
        print(f"   找到图表标题: {text[:80]}...")

if figure_captions_found == 0:
    print("   ⚠️ 未找到图表标题，需要手动添加")

# 4. 保存文档
output_file = "main_paper_jfds_READY_v11_final.docx"
doc.save(output_file)

print(f"\n✅ 修复完成！")
print(f"保存为: {output_file}")

# 5. 验证修复
print("\n=== 验证修复 ===")
doc2 = docx.Document(output_file)

# 重新统计
table_refs = set()
table_captions = set()
for para in doc2.paragraphs:
    text = para.text.strip()
    refs = re.findall(r'Table\s+(\d+)', text, re.IGNORECASE)
    table_refs.update(refs)
    if re.match(r'^Table\s+\d+', text, re.IGNORECASE):
        caption_num = re.findall(r'Table\s+(\d+)', text, re.IGNORECASE)
        table_captions.update(caption_num)

print(f"表格引用: {sorted(table_refs)}")
print(f"表格标题: {sorted(table_captions)}")

# 检查附录D重复
appendix_d_count = 0
for para in doc2.paragraphs:
    if 'Appendix D: Reproducibility' in para.text:
        appendix_d_count += 1

print(f"附录D出现次数: {appendix_d_count}")

