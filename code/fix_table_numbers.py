import docx
from docx.shared import Pt
import re

# 加载文档
doc = docx.Document("/Volumes/agents/codex/paper2-LSTM量化金融/archive_20260710_1046_chatgpt_review/JFDS_Submission_Package_20260712/paper/论文_公式恢复与图表美化_NUMBER_FIXED.docx")

print("=== 修正表格中的数字 ===\n")

# 修正表格中的234
fixed_count = 0
for i, table in enumerate(doc.tables):
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            original_text = cell.text
            if '234' in original_text:
                # 将234替换为228
                new_text = original_text.replace('234', '228')
                # 清除并重新设置单元格内容
                for para in cell.paragraphs:
                    para.clear()
                    para.add_run(new_text)
                fixed_count += 1
                print(f"  修正 Table {i+1}, Row {row_idx}, Cell {cell_idx}:")
                print(f"    原文: {original_text[:50]}...")
                print(f"    修正: {new_text[:50]}...")

# 保存文档
output_file = "/Volumes/agents/codex/paper2-LSTM量化金融/archive_20260710_1046_chatgpt_review/JFDS_Submission_Package_20260712/paper/论文_公式恢复与图表美化_FINAL.docx"
doc.save(output_file)

print(f"\n✅ 表格数字修正完成！")
print(f"修正单元格数: {fixed_count}")
print(f"保存为: {output_file}")

