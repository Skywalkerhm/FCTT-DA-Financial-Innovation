import docx
from docx.shared import Pt

def read_markdown_file(filepath):
    """Read a markdown file and return its content."""
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()

def replace_appendix_d(doc, new_content):
    """Replace Appendix D placeholder with actual content."""
    lines = new_content.strip().split('\n')
    
    # Find Appendix D placeholder
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if "[SHA-256 checksums, protocol JSON, environment specs]" in text:
            # Replace this placeholder with new content
            # First, clear the current paragraph
            para.clear()
            
            # Add the first line
            first_line = lines[0].strip()
            if first_line.startswith('## '):
                para.add_run(first_line[3:])
                para.style = doc.styles['Heading 1']
            else:
                para.add_run(first_line)
            
            # Add remaining lines after this paragraph
            last_para = para
            for line in lines[1:]:
                line = line.strip()
                if not line:
                    continue
                
                new_para = doc.add_paragraph()
                
                if line.startswith('### '):
                    new_para.style = doc.styles['Heading 2']
                    new_para.add_run(line[4:])
                elif line.startswith('**Table'):
                    run = new_para.add_run(line)
                    run.bold = True
                elif line.startswith('```'):
                    # Code block - skip for now
                    continue
                elif line.startswith('- '):
                    new_para.style = doc.styles['List Bullet']
                    new_para.add_run(line[2:])
                else:
                    new_para.add_run(line)
                
                last_para._element.addnext(new_para._element)
                last_para = new_para
            
            print(f"✅ 替换了附录D占位符 (段落 {i})")
            return True
    
    print("❌ 未找到附录D占位符")
    return False

# Load the document
doc = docx.Document("main_paper_jfds_READY_v8_transaction_costs.docx")

print("=== 更新附录D: Reproducibility ===\n")

# Read the new appendix content
new_content = read_markdown_file("appendix_d_reproducibility.md")

# Replace the placeholder
replace_appendix_d(doc, new_content)

# Save the document
output_file = "main_paper_jfds_READY_v9_reproducibility.docx"
doc.save(output_file)

print(f"\n✅ 附录D更新完成！")
print(f"保存为: {output_file}")

# Count total words
import re
total_words = sum(len(re.findall(r'\w+', para.text)) for para in doc.paragraphs)
print(f"文档总字数: {total_words}")

