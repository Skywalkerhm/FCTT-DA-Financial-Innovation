import docx
from docx.shared import Pt
import re

# 加载文档
doc = docx.Document("/Volumes/agents/codex/paper2-LSTM量化金融/archive_20260710_1046_chatgpt_review/JFDS_Submission_Package_20260712/paper/论文_公式恢复与图表美化_RHETORIC.docx")

print("=== 修复数字一致性问题 ===\n")

# 定义替换规则
replacements = [
    # Abstract中的234
    ('Across 234 source-target pairs', 'Across 228 source-target pairs'),
    
    # 其他可能的234
    ('234 source-target pairs', '228 source-target pairs'),
    ('234 pairs', '228 pairs'),
    ('of 234', 'of 228'),
    ('for 234', 'for 228'),
    
    # 39 targets的解释
    ('39 target assets', '39 assets (38 per source after excluding self)'),
    
    # Table 4中的234
    ('234 pairs', '228 pairs'),
]

# 修正文档
fixed_count = 0
fixed_paragraphs = []

for i, para in enumerate(doc.paragraphs):
    original_text = para.text
    fixed_text = original_text
    
    # 应用替换
    for old, new in replacements:
        if old in fixed_text:
            fixed_text = fixed_text.replace(old, new)
            fixed_count += 1
    
    # 更新段落
    if fixed_text != original_text:
        para.clear()
        para.add_run(fixed_text)
        fixed_paragraphs.append((i, original_text[:60], fixed_text[:60]))
        print(f"  修正段落 {i}:")
        print(f"    原文: {original_text[:60]}...")
        print(f"    修正: {fixed_text[:60]}...")

# 同时修正表格数据
print("\n4. 检查表格数据:")
for i, table in enumerate(doc.tables):
    for row in table.rows:
        for cell in row.cells:
            if '234' in cell.text:
                print(f"   Table {i+1}: 找到234 - {cell.text[:50]}...")

# 保存文档
output_file = "/Volumes/agents/codex/paper2-LSTM量化金融/archive_20260710_1046_chatgpt_review/JFDS_Submission_Package_20260712/paper/论文_公式恢复与图表美化_NUMBER_FIXED.docx"
doc.save(output_file)

print(f"\n✅ 数字一致性修复完成！")
print(f"修正段落数: {len(fixed_paragraphs)}")
print(f"保存为: {output_file}")

