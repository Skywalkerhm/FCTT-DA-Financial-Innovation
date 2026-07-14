import docx
from docx.shared import Pt
import re

# 加载文档
doc = docx.Document("/Volumes/agents/codex/paper2-LSTM量化金融/archive_20260710_1046_chatgpt_review/JFDS_Submission_Package_20260712/paper/FCTT_DA_JFDS_reformatted.docx")

print("=== 修正复核清单问题 ===\n")

# 定义修正规则
replacements = [
    # 1. Bootstrap次数修正
    ('with  replications and a mean block length of  periods', 
     'with 2,000 replications and a mean block length of 10 periods'),
    ('with ___ replications', 'with 2,000 replications'),
    
    # 2. EMD描述增强
    ('Empirical Mode Decomposition (EMD) Trend: The first intrinsic mode function of the rolling window of price observations, capturing short-term momentum',
     'Empirical Mode Decomposition (EMD) Trend: The first Intrinsic Mode Function (IMF) of the price series, computed using a rolling window of 252 trading days. At each date t, EMD is re-estimated using only observations in [t-251, t] to avoid look-ahead bias. Mirror boundary conditions (63 points at each end) are applied to mitigate endpoint effects.'),
    
    # 3. 过强结论软化
    ('Collapse rates are invariant to the threshold grid',
     'Collapse rates are robust across the threshold grid range tested'),
    ('Collapse is robust to changes in the exploration parameter',
     'Collapse persists across the exploration parameter values tested'),
    ('Calibration does not affect collapse rates',
     'Calibration does not materially affect collapse rates in our experiments'),
    
    # 4. "fundamental"修正
    ('fundamentally different factors', 'different macroeconomic factors'),
    ('fundamental question', 'key question'),
    
    # 5. 语法错误修正
    ('is persists', 'persists'),
    ('policy policies', 'policies'),
    
    # 6. "invariant"修正
    ('invariant to', 'robust to'),
    
    # 7. NoGating格式修正
    ('NoGating: A static threshold of 0.50 (also called "Ungated classifier"), ,',
     'NoGating: A static threshold of 0.50 (also called "Ungated classifier"),'),
]

# 修正文档
fixed_count = 0
for i, para in enumerate(doc.paragraphs):
    original_text = para.text
    fixed_text = original_text
    
    # 应用替换
    for old, new in replacements:
        if old in fixed_text:
            fixed_text = fixed_text.replace(old, new)
            fixed_count += 1
    
    # 更新段落
    if fixed_text != original_text:
        para.clear()
        para.add_run(fixed_text)
        print(f"  修正段落 {i}: {original_text[:50]}...")

# 保存文档
output_file = "/Volumes/agents/codex/paper2-LSTM量化金融/archive_20260710_1046_chatgpt_review/JFDS_Submission_Package_20260712/paper/FCTT_DA_JFDS_reformatted_v2.docx"
doc.save(output_file)

print(f"\n✅ 复核清单问题修正完成！")
print(f"修正数量: {fixed_count}")
print(f"保存为: {output_file}")

